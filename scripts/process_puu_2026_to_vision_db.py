#!/usr/bin/env python3
"""
Process PUU 2026 Files to Vision Database

Script untuk memproses file PDF dan DOCX dari folder OneDrive_PUU/PUU_2026
menggunakan sistem hybrid storage dengan confidence-based filtering.

Author: AI Assistant
Date: 2026-03-03
"""

import asyncio
import sys
import os
import hashlib
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional

# Add project root to path
sys.path.insert(0, '/home/aseps/MCP/mcp-unified')

from core.vision_config import get_config, classify_document_type, calculate_content_quality
from memory.vision_repository import save_vision_result, check_duplicate
from core.vision_config import ProcessingResult

# Document processing imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️  PyMuPDF not available. PDF processing disabled.")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    print("⚠️  python-docx not available. DOCX processing disabled.")


class PUU2026Processor:
    """Processor untuk file PUU 2026 dengan hybrid storage"""
    
    def __init__(self, input_dir: str = "OneDrive_PUU/PUU_2026", 
                 namespace: str = "PUU_2026"):
        self.input_dir = Path(input_dir)
        self.namespace = namespace
        self.config = get_config()
        
        # Statistics
        self.stats = {
            'total_files': 0,
            'processed': 0,
            'saved_to_sql': 0,
            'saved_to_ltm': 0,
            'rejected': 0,
            'skipped': 0,
            'errors': 0
        }
        
    def calculate_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception as e:
            print(f"❌ Error calculating hash: {e}")
            return ""
    
    def extract_text_from_pdf(self, pdf_path: Path) -> tuple:
        """Extract text from PDF file"""
        if not PYMUPDF_AVAILABLE:
            return "", 0
        
        try:
            doc = fitz.open(pdf_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            page_count = len(text_parts)
            
            return full_text, page_count
            
        except Exception as e:
            print(f"❌ Error extracting PDF {pdf_path}: {e}")
            return "", 0
    
    def extract_text_from_docx(self, docx_path: Path) -> tuple:
        """Extract text from DOCX file"""
        if not PYTHON_DOCX_AVAILABLE:
            return "", 0
        
        try:
            doc = Document(docx_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            
            return full_text, len(text_parts)
            
        except Exception as e:
            print(f"❌ Error extracting DOCX {docx_path}: {e}")
            return "", 0
    
    def calculate_confidence(self, text: str, file_type: str) -> float:
        """Calculate confidence score based on content quality"""
        base_confidence = 0.6  # Base confidence for extracted text
        
        # Content length factor
        text_length = len(text)
        if text_length > 5000:
            base_confidence += 0.15
        elif text_length > 1000:
            base_confidence += 0.10
        elif text_length > 500:
            base_confidence += 0.05
        
        # Content quality metrics
        quality = calculate_content_quality(text)
        base_confidence += quality['quality_score'] * 0.2
        
        # File type bonus
        if file_type == 'pdf':
            base_confidence += 0.05  # PDFs are more reliable
        
        # Document classification bonus
        doc_type, boost = classify_document_type(text)
        base_confidence += boost
        
        return min(1.0, max(0.0, base_confidence))
    
    async def process_file(self, file_path: Path) -> Dict[str, Any]:
        """Process single file"""
        print(f"\n📄 Processing: {file_path.name}")
        
        # Calculate file hash
        file_hash = self.calculate_file_hash(file_path)
        
        # Check for duplicates
        dup_check = await check_duplicate(file_hash, self.namespace)
        if dup_check.get('exists'):
            print(f"   ⚠️  Duplicate found (ID: {dup_check.get('id')})")
            self.stats['skipped'] += 1
            return {'success': False, 'reason': 'duplicate'}
        
        # Extract text based on file type
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.pdf':
            text, page_count = self.extract_text_from_pdf(file_path)
            mime_type = 'application/pdf'
        elif file_ext in ['.docx', '.doc']:
            text, line_count = self.extract_text_from_docx(file_path)
            page_count = line_count
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            print(f"   ⚠️  Unsupported file type: {file_ext}")
            return {'success': False, 'reason': 'unsupported_type'}
        
        if not text:
            print(f"   ⚠️  No text extracted")
            self.stats['errors'] += 1
            return {'success': False, 'reason': 'no_text'}
        
        print(f"   📊 Extracted {len(text)} characters ({page_count} pages/lines)")
        
        # Calculate confidence
        confidence = self.calculate_confidence(text, file_ext.replace('.', ''))
        print(f"   🎯 Confidence: {confidence:.2f}")
        
        # Classify document
        doc_type, _ = classify_document_type(text)
        print(f"   🏷️  Document type: {doc_type}")
        
        # Get storage decision
        storage_decision = self.config.get_storage_decision(confidence)
        print(f"   💾 Storage decision: {storage_decision}")
        
        # Prepare entities
        entities = {
            "dates": self._extract_dates(text),
            "amounts": self._extract_amounts(text),
            "document_type": doc_type,
            "page_count": page_count,
            "char_count": len(text)
        }
        
        # Create processing result
        result = ProcessingResult(
            file_name=file_path.name,
            file_path=str(file_path),
            file_hash=file_hash,
            file_size_bytes=file_path.stat().st_size,
            mime_type=mime_type,
            extracted_text=text,
            confidence_score=confidence,
            processing_method="manual",  # Valid values: 'vision', 'ocr', 'hybrid', 'manual'
            model_used="pymupdf" if file_ext == '.pdf' else "python-docx",
            processing_time_ms=0,  # Will be calculated
            document_type=doc_type,
            status="success",
            extracted_entities=entities,
            processing_metadata={
                "source": "PUU_2026_batch",
                "file_ext": file_ext,
                "processed_at": datetime.now().isoformat()
            },
            namespace=self.namespace,
            tenant_id="default",
            ltm_key=""
        )
        
        # Save to database
        save_result = await save_vision_result(result, self.config)
        
        if save_result['saved_to_sql']:
            self.stats['saved_to_sql'] += 1
            print(f"   ✅ Saved to SQL (ID: {save_result.get('id')})")
        
        if save_result.get('storage_decision') == 'sql+ltm':
            self.stats['saved_to_ltm'] += 1
        
        if save_result['storage_decision'] == 'reject':
            self.stats['rejected'] += 1
            print(f"   ❌ Rejected (low confidence)")
        
        self.stats['processed'] += 1
        
        return {
            'success': True,
            'confidence': confidence,
            'storage_decision': storage_decision,
            'document_type': doc_type
        }
    
    def _extract_dates(self, text: str) -> List[str]:
        """Extract dates from text"""
        import re
        patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',
        ]
        dates = []
        for pattern in patterns:
            dates.extend(re.findall(pattern, text))
        return dates[:10]  # Limit to 10 dates
    
    def _extract_amounts(self, text: str) -> List[str]:
        """Extract monetary amounts"""
        import re
        patterns = [
            r'Rp\s*[\d.,]+',
            r'IDR\s*[\d.,]+',
        ]
        amounts = []
        for pattern in patterns:
            amounts.extend(re.findall(pattern, text))
        return amounts[:10]
    
    async def process_directory(self):
        """Process all files in directory"""
        print("="*70)
        print("🚀 PUU 2026 File Processor with Hybrid Storage")
        print("="*70)
        print(f"📁 Directory: {self.input_dir}")
        print(f"🏷️  Namespace: {self.namespace}")
        print(f"⚙️  SQL Threshold: {self.config.get_threshold('high')}")
        print(f"⚙️  LTM Threshold: {self.config.get_threshold('medium')}")
        print("="*70)
        
        # Find all PDF and DOCX files
        pdf_files = list(self.input_dir.glob("*.pdf"))
        docx_files = list(self.input_dir.glob("*.docx"))
        doc_files = list(self.input_dir.glob("*.doc"))
        
        all_files = pdf_files + docx_files + doc_files
        self.stats['total_files'] = len(all_files)
        
        print(f"\n📊 Found {len(all_files)} files:")
        print(f"   📄 PDF: {len(pdf_files)}")
        print(f"   📝 DOCX: {len(docx_files)}")
        print(f"   📝 DOC: {len(doc_files)}")
        
        # Process files
        for i, file_path in enumerate(all_files, 1):
            print(f"\n[{i}/{len(all_files)}]", end="")
            try:
                await self.process_file(file_path)
            except Exception as e:
                print(f"   ❌ Error: {e}")
                self.stats['errors'] += 1
        
        # Print summary
        self.print_summary()
    
    def print_summary(self):
        """Print processing summary"""
        print("\n" + "="*70)
        print("📊 PROCESSING SUMMARY")
        print("="*70)
        print(f"📁 Total files: {self.stats['total_files']}")
        print(f"✅ Processed: {self.stats['processed']}")
        print(f"💾 Saved to SQL: {self.stats['saved_to_sql']}")
        print(f"🧠 Saved to LTM: {self.stats['saved_to_ltm']}")
        print(f"❌ Rejected: {self.stats['rejected']}")
        print(f"⏭️  Skipped (duplicates): {self.stats['skipped']}")
        print(f"💥 Errors: {self.stats['errors']}")
        print("="*70)
        
        if self.stats['total_files'] > 0:
            success_rate = (self.stats['processed'] / self.stats['total_files']) * 100
            print(f"📈 Success rate: {success_rate:.1f}%")
        
        print("\n✨ Processing complete!")


async def main():
    """Main entry point"""
    processor = PUU2026Processor()
    await processor.process_directory()


if __name__ == "__main__":
    asyncio.run(main())
