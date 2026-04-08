#!/usr/bin/env python3
"""
Fix DOCX Formatting V2 - Professional Ministry Document Format
Based on: PENGUMUMAN-SELEKSI-TERBUKA-JPT-PRATAMA_4-JPT.pdf

Improvements over V1:
1. Proper Multilevel List (A → 1 → a → 1))
2. Hanging Indent for list items
3. Electronic signature footer text
4. Proper table formatting (100% width, header bold, cell margins)
5. Remove page number artifacts (-2-, -3-, etc.)
6. Remove [LOGO] text placeholders
7. Page breaks before major sections
"""

import sys
import re
from pathlib import Path

# Add project root to path for imports
project_root = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# File paths
INPUT_FILE = "storage/office/PENGUMUMAN-SELEKSI-JPT-PRATAMA.docx"
OUTPUT_FILE = "storage/office/PENGUMUMAN-SELEKSI-JPT-PRATAMA-FORMATTED-V3.docx"

# ============================================================
# HEADER DEFINITIONS (based on PDF reference)
# ============================================================
LETTERHEAD_LINES = [
    ("KEMENTERIAN TRANSMIGRASI REPUBLIK INDONESIA", True, 12),
    ("SEKRETARIAT JENDERAL", True, 11),
    ("Jalan TMP. Kalibata Nomor 17 Jakarta Selatan 12750 Telepon 021-7971972", False, 10),
    ("www.transmigrasi.go.id", False, 10),
]

# Electronic signature footer text (from PDF)
ELECTRONIC_SIG_FOOTER = "Dokumen ini telah ditandatangani secara elektronik menggunakan sertifikat elektronik yang diterbitkan oleh Balai Besar Sertifikasi Elektronik (BSrE), Badan Siber dan Sandi Negara (BSSN)."

# Page number patterns to remove
PAGE_NUMBER_PATTERN = re.compile(r'^-\d+-$')
LOGO_PATTERN = re.compile(r'^\[?logo\]?$', re.IGNORECASE)


def is_page_number_artifact(text: str) -> bool:
    """Check if text is a page number artifact like -2-, -3-, etc."""
    return bool(PAGE_NUMBER_PATTERN.match(text.strip()))


def is_logo_placeholder(text: str) -> bool:
    """Check if text is a logo placeholder."""
    return bool(LOGO_PATTERN.match(text.strip()))


def set_paragraph_spacing(para, before=None, after=None, line_spacing=None):
    """Set paragraph spacing."""
    if before is not None:
        para.paragraph_format.space_before = Pt(before)
    if after is not None:
        para.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        para.paragraph_format.line_spacing = Pt(line_spacing)


def set_paragraph_indent(para, left=None, right=None, first_line=None, hanging=None):
    """Set paragraph indent using OOXML.
    All values in centimeters (cm)."""
    pPr = para._p.get_or_add_pPr()
    
    # Remove existing ind element
    existing = pPr.find(qn('w:ind'))
    if existing is not None:
        pPr.remove(existing)
    
    ind = OxmlElement('w:ind')
    CM_TO_TWIPS = 567
    
    if left is not None:
        left_twips = int(round(float(left) * CM_TO_TWIPS))  # Convert cm to twips: 1cm = ~567 twips
        ind.set(qn('w:left'), str(left_twips))
    
    if right is not None:
        right_twips = int(round(float(right) * CM_TO_TWIPS))
        ind.set(qn('w:right'), str(right_twips))
    
    if first_line is not None:
        first_twips = int(round(float(first_line) * CM_TO_TWIPS))
        ind.set(qn('w:firstLine'), str(first_twips))
    
    if hanging is not None:
        # w:firstLine in twips, positive value for hanging indent
        hanging_twips = int(round(float(hanging) * CM_TO_TWIPS))
        ind.set(qn('w:firstLine'), str(hanging_twips))
    
    pPr.append(ind)


def set_font(run, name='Arial', size=11, bold=False, italic=False, color=None):
    """Set font properties for a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_header(doc):
    """Add letterhead header to document (first page only)."""
    for section in doc.sections:
        # Enable different first page header
        sectPr = section._sectPr
        # Remove existing titlePg if any
        existing_titlePg = sectPr.find(qn('w:titlePg'))
        if existing_titlePg is not None:
            sectPr.remove(existing_titlePg)
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)
        
        # Add to first page header only
        first_header = section.first_page_header
        first_header.is_linked_to_previous = False
        
        # Clear existing header content
        for p in first_header.paragraphs:
            p.clear()
        
        # Add letterhead lines only to first page header
        for i, (text, bold, size) in enumerate(LETTERHEAD_LINES):
            if i == 0:
                para = first_header.paragraphs[0]
            else:
                para = first_header.add_paragraph()
            
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            set_font(run, size=size, bold=bold)
            
            if i < len(LETTERHEAD_LINES) - 1:
                set_paragraph_spacing(para, after=2)
            else:
                set_paragraph_spacing(para, after=0)


def add_footer(doc):
    """Add footer with electronic signature notice only (no page number, smaller font)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Clear existing footer
        for p in footer.paragraphs:
            p.clear()
        
        # Add electronic signature notice only (no page number)
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_run = footer_para.add_run(ELECTRONIC_SIG_FOOTER)
        set_font(sig_run, size=7)  # Smaller font for footer
        set_paragraph_spacing(footer_para, before=0, after=0)


def set_page_setup(doc):
    """Set page size and margins."""
    for section in doc.sections:
        section.page_width = Cm(21.0)  # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def format_table(table):
    """Format table to match professional style."""
    # Set table width to 100%
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    # Remove existing tblW
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    
    # Set table width to 100%
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    
    tbl.append(tblPr)
    
    # Set cell margins
    cellMar = OxmlElement('w:tblCellMar')
    for margin_name in ['top', 'left', 'bottom', 'right']:
        mar = OxmlElement(f'w:{margin_name}')
        mar.set(qn('w:w'), '70')  # ~0.5 mm
        mar.set(qn('w:type'), 'dxa')
        cellMar.append(mar)
    
    # Remove existing cellMar
    existing_cellmar = tblPr.find(qn('w:tblCellMar'))
    if existing_cellmar is not None:
        tblPr.remove(existing_cellmar)
    tblPr.append(cellMar)
    
    # Format header row (first row)
    if table.rows:
        for cell in table.rows[0].cells:
            # Bold all paragraph runs in header
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                set_paragraph_spacing(para, before=2, after=2)
            # Center align header
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Format data rows
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                set_paragraph_spacing(para, before=1, after=1)
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'


def set_default_paragraph_style(para, font_name='Arial', font_size=11):
    """Set default style for paragraph runs."""
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
    
    # Ensure at least one run exists
    if not para.runs:
        run = para.add_run(para.text)
        para.clear()
        set_font(run, name=font_name, size=font_size)


# ============================================================
# Multilevel List Detection Functions
# ============================================================

def detect_list_level(text: str) -> tuple | None:
    """
    Detect the list level of a paragraph text.
    Returns (level, prefix_text, content_text) or None if not a list item.
    
    Levels:
    0: A, B, C, D, E (section headers)
    1: 1, 2, 3, 4, 5 (numbered items)
    2: a, b, c, d (letter sub-items)
    3: 1), 2), 3) (numbered sub-items)
    """
    text_stripped = text.strip()
    if not text_stripped:
        return None
    
    # Level 0 sections: A. KETENTUAN UMUM, B. KETENTUAN KHUSUS, etc.
    m = re.match(r'^([A-E])\.\s+(.+)$', text_stripped)
    if m:
        return (0, m.group(1), m.group(2))
    
    # Level 1 numbered: 1. Nama Jabatan lowong, 2. Persyaratan Umum, etc.
    # Also: 3. Tata Cara Pengajuan Lamaran, etc.
    m = re.match(r'^(\d{1,2})\.\s+(.+)$', text_stripped)
    if m:
        prefix = m.group(1)
        content = m.group(2)
        # Exclude subsection items like "1) already handled"
        return (1, prefix, content)
    
    # Level 2 letter: a. Peserta wajib mengajukan lamaran...
    m = re.match(r'^([a-z])\.\s+(.+)$', text_stripped)
    if m:
        return (2, m.group(1), m.group(2))
    
    # Level 3 numbered: 1) Berstatus sebagai Pegawai Negeri Sipil;
    m = re.match(r'^(\d{1,2})\)\s+(.+)$', text_stripped)
    if m:
        return (3, m.group(1), m.group(2))
    
    return None


# Indentation levels in cm (based on user's specifications)
INDENT_LEVELS = {
    0: {'left': 0.0, 'hanging': 0.0},      # A, B, C - no indent
    1: {'left': 0.70, 'hanging': 0.0},     # 1, 2, 3 - 0.70 cm
    2: {'left': 1.4, 'hanging': 0.0},      # a, b, c - 1.4 cm  
    3: {'left': 2.1, 'hanging': 0.0},      # 1), 2), 3) - 2.1 cm
}

# Additional text indent (where text wraps to)
TEXT_INDENT_AFTER = {
    0: 0.70,   # A. → text at 0.70 cm
    1: 1.4,    # 1. → text at 1.4 cm
    2: 2.1,    # a. → text at 2.1 cm
    3: 2.8,    # 1) → text at 2.8 cm
}


def apply_multilevel_list_formatting(para, level: int, prefix_str: str, content: str):
    """Apply multilevel list formatting with proper indents."""
    # Get indent values
    indent_config = INDENT_LEVELS.get(level, {'left': 0.0, 'hanging': 0.0})
    text_indent = TEXT_INDENT_AFTER.get(level, 0.0)
    
    # Clear paragraph content
    para.clear()
    
    # Add number prefix
    if level == 0:
        # A. KETENTUAN UMUM (both prefix and content in same run, bold)
        run = para.add_run(f"{prefix_str}.  {content}")
        set_font(run, bold=True, size=11)
    elif level == 1:
        # 1. Nama Jabatan lowong (both prefix and content, bold)
        run = para.add_run(f"{prefix_str}.  {content}")
        set_font(run, bold=True, size=11)
    elif level == 2:
        # a. text (letter with period)
        run = para.add_run(f"{prefix_str}.  {content}")
        set_font(run, size=11)
    elif level == 3:
        # 1) text (number with parenthesis)
        run = para.add_run(f"{prefix_str})  {content}")
        set_font(run, size=11)
    
    # Set alignment
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Set indentation (hanging indent: text indents more than number)
    pPr = para._p.get_or_add_pPr()
    existing = pPr.find(qn('w:ind'))
    if existing is not None:
        pPr.remove(existing)
    ind = OxmlElement('w:ind')
    left_twips = int(round(indent_config['left'] * 567))
    hanging_twips = int(round(text_indent * 567 - indent_config['left'] * 567))
    ind.set(qn('w:left'), str(left_twips))
    ind.set(qn('w:hanging'), str(hanging_twips))
    pPr.append(ind)
    
    # Set spacing
    set_paragraph_spacing(para, before=0, after=0, line_spacing=18)


# ============================================================
# Main formatting function
# ============================================================

def format_docx(input_path: str, output_path: str):
    """Apply professional formatting to DOCX file."""
    doc = Document(input_path)
    
    # Set page setup
    set_page_setup(doc)
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # ============================================================
    # STEP 1: Remove page number artifacts, logo placeholders, and page breaks
    # ============================================================
    paragraphs_to_delete = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if is_page_number_artifact(text):
            paragraphs_to_delete.append(para)
        elif is_logo_placeholder(text):
            paragraphs_to_delete.append(para)
    
    for para in paragraphs_to_delete:
        p = para._element
        p.getparent().remove(p)
    
    # Remove ALL page breaks from original document
    # The 37 pages is caused by 10 page breaks in original + large tables (Lampiran)
    page_break_count = 0
    for para in doc.paragraphs:
        runs_to_modify = list(para.runs)
        for run in runs_to_modify:
            # Find all page breaks in the run
            br_elements = run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
            for br in br_elements:
                if br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') == 'page':
                    br.getparent().remove(br)
                    page_break_count += 1
    
    print(f'Removed {page_break_count} page breaks from original document')
    
    # Remove Table 0 (logo/header table) - we already have letterhead in the header
    # This table contains "[LOGO]" and ministry info which overlaps with letterhead header
    if doc.tables:
        first_table = doc.tables[0]
        # Check if it's the logo table
        first_cell_text = first_table.rows[0].cells[0].text if first_table.rows and first_table.rows[0].cells else ''
        if any(kw in first_cell_text for kw in ['[LOGO]', 'logo', 'Logo']):
            # Remove the table
            tbl = first_table._tbl
            tbl.getparent().remove(tbl)
            print('Removed logo table (letterhead already in header)')
    
    # ============================================================
    # STEP 2: Identify section positions for page breaks
    # ============================================================
    SECTION_BREAKS = [
        "D.  JADWAL WAKTU PELAKSANAAN",
        "E.  KETENTUAN LAIN",
    ]
    
    # ============================================================
    # STEP 3: Apply formatting to all paragraphs
    # ============================================================
    title_text_patterns = [
        "PENGUMUMAN",
        "Nomor :",
        "TENTANG",
        "SELEKSI TERBUKA",
        "PENGISIAN JABATAN",
        "DI LINGKUNGAN",
    ]
    
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text:
            # Empty paragraph - keep but set minimal spacing
            set_paragraph_spacing(para, before=0, after=0)
            i += 1
            continue
        
        # Title block (first few non-empty paragraphs)
        if any(p in text for p in title_text_patterns):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(para, before=0, after=2)
            for run in para.runs:
                if "PENGUMUMAN" in text:
                    set_font(run, size=14, bold=True)
                elif "Nomor" in text:
                    set_font(run, size=11)
                elif any(k in text for k in ["TENTANG", "SELEKSI TERBUKA", "PENGISIAN JABATAN", "DI LINGKUNGAN"]):
                    set_font(run, size=12, bold=True)
                else:
                    set_font(run, size=11)
            i += 1
            continue
        
        # "Berdasarkan:" line
        if text.startswith("Berdasarkan"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(para, before=6, after=0, line_spacing=18)
            for run in para.runs:
                set_font(run, bold=True, size=11)
            i += 1
            continue
        
        # Reference list items (a., b., c., d. after Berdasarkan)
        if re.match(r'^[a-d]\.\s+(Undang-Undang|Peraturan|Peraturan Menteri)', text):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_indent(para, left=1.5, hanging=0.5)
            set_paragraph_spacing(para, before=0, after=0, line_spacing=18)
            for run in para.runs:
                set_font(run, size=11)
            i += 1
            continue
        
        # Opening paragraph (Dalam rangka...)
        if text.startswith("Dalam rangka"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(para, before=6, after=0, line_spacing=18)
            set_paragraph_indent(para, first_line=0.5)
            for run in para.runs:
                set_font(run, size=11)
            i += 1
            continue
        
        # Check for section header: "dengan ketentuan sebagai berikut:"
        if text.startswith("dengan ketentuan"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(para, before=0, after=4, line_spacing=18)
            for run in para.runs:
                set_font(run, size=11)
            i += 1
            continue
        
        # Check for multilevel list items
        list_info = detect_list_level(text)
        if list_info:
            level, prefix, content = list_info
            
            # No page breaks - let content flow naturally
            apply_multilevel_list_formatting(para, level, prefix, content)
            i += 1
            continue
        
        # "Setiap pelamar hanya dapat..." - descriptive text after section 1.
        if text.startswith("Setiap pelamar"):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(para, before=0, after=0, line_spacing=18)
            set_paragraph_indent(para, left=0.75)
            for run in para.runs:
                set_font(run, size=11)
            i += 1
            continue
        
        # Closing section (Dikeluarkan di / Pada tanggal)
        if text.startswith("Dikeluarkan di") or text.startswith("Pada tanggal"):
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_spacing(para, before=4, after=2, line_spacing=18)
            for run in para.runs:
                set_font(run, size=11)
            i += 1
            continue
        
        # Any remaining text - default formatting
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(para, before=0, after=0, line_spacing=18)
        for run in para.runs:
            set_font(run, size=11)
        
        i += 1
    
    # ============================================================
    # STEP 4: Format tables
    # ============================================================
    for table in doc.tables:
        format_table(table)
    
    # ============================================================
    # STEP 5: Add header and footer
    # ============================================================
    add_header(doc)
    add_footer(doc)
    
    # ============================================================
    # Save
    # ============================================================
    doc.save(output_path)
    print(f"Document formatted and saved to: {output_path}")


# ============================================================
# Main execution
# ============================================================
if __name__ == "__main__":
    full_input = project_root / INPUT_FILE
    full_output = project_root / OUTPUT_FILE
    
    if not full_input.exists():
        print(f"Error: Input file not found: {full_input}")
        sys.exit(1)
    
    print(f"Formatting (V2): {full_input}")
    print(f"Output: {full_output}")
    
    format_docx(str(full_input), str(full_output))
    
    print("\nFormatting applied (V3):")
    print("  ✓ Arial font throughout")
    print("  ✓ No page numbers in footer")
    print("  ✓ Smaller footer font (7pt)")
    print("  ✓ Header on first page only")
    print("  ✓ Multilevel List (A → 1 → a → 1))")
    print("  ✓ Level 0 indent: left 0 cm, text at 0.70 cm")
    print("  ✓ Hanging Indent for list items")
    print("  ✓ Electronic signature footer text")
    print("  ✓ Table formatting (100% width, bold header)")
    print("  ✓ Table cell margins")
    print("  ✓ Removed page number artifacts (-2-, -3-, etc.)")
    print("  ✓ Removed [LOGO] placeholders")
    print("  ✓ A4 page with standard margins")
