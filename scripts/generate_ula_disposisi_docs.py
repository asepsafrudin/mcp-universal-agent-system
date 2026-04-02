#!/usr/bin/env python3
"""
Generator Dokumen Disposisi ULA (Surat Luar 2026)
==================================================
Menghasilkan dokumen disposisi dari data surat_keluar_ula dan lembar_disposisi_ula.

Format: DOCX dengan header disposisi sederhana.

Jalankan:
  python3 scripts/generate_ula_disposisi_docs.py
  python3 scripts/generate_ula_disposisi_docs.py --dry-run
  python3 scripts/generate_ula_disposisi_docs.py --limit 50
"""
import io
import os
import sys
import json
import argparse
import logging
import re
from datetime import datetime, date
from typing import Any, Dict, Optional

import psycopg
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = "/home/aseps/MCP/mcp-unified"
sys.path.insert(0, PROJECT_ROOT)
from core.secrets import load_runtime_secrets
load_runtime_secrets()

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("generate_ula_disposisi")

# Configuration
LOCAL_DOCS_DIR = "/home/aseps/MCP/data/storage/disposisi_docs_ula"
os.makedirs(LOCAL_DOCS_DIR, exist_ok=True)


def fmt_date(d: Any) -> str:
    """Format tanggal Indonesia."""
    BULAN = {1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni",
             7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober", 11: "November", 12: "Desember"}
    if isinstance(d, (datetime, date)):
        return f"{d.day:02d} {BULAN.get(d.month, '')} {d.year}"
    if isinstance(d, str):
        try:
            dt = datetime.fromisoformat(d.replace('Z', '+00:00'))
            return fmt_date(dt)
        except (ValueError, TypeError):
            pass
    return str(d) if d else ""


def create_disposisi_document(data: Dict[str, Any]) -> bytes:
    """Create a simple disposisi DOCX document."""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("LEMBAR DISPOSISI")
    run.bold = True
    run.font.size = Pt(14)
    
    # Separator line
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("_" * 50)
    run.font.size = Pt(10)
    
    # Data table
    def add_row(label: str, value: str, bold_label: bool = True):
        """Add a row of data."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if bold_label:
            run_label = p.add_run(label)
            run_label.bold = True
            p.add_run(f": {value or '-'}")
        else:
            p.add_run(f"{label}: {value or '-'}")
    
    add_row("Nomor Agenda", data.get("agenda_ula"))
    add_row("Surat Dari", data.get("surat_dari"))
    add_row("Nomor Surat", data.get("nomor_surat"))
    add_row("Tanggal Surat", fmt_date(data.get("tgl_surat")))
    add_row("Tanggal Diterima ULA", fmt_date(data.get("tgl_diterima_ula")))
    add_row("Perihal", data.get("perihal"))
    add_row("Arahan Menteri", data.get("arahan_menteri"))
    add_row("Arahan Sekjen", data.get("arahan_sekjen"))
    add_row("Status Mailmerge", data.get("status_mailmerge"))
    
    # Disposisi section
    if data.get("disposisi_content"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("DISPOSISI / ARAHAN")
        run.bold = True
        run.font.size = Pt(12)
        
        p = doc.add_paragraph(data.get("disposisi_content", ""))
        p.paragraph_format.space_before = Pt(6)
    
    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Dicetak oleh: Sistem MCP ULA\n")
    run.font.size = Pt(9)
    p.add_run(f"Tanggal: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    for r in p.runs:
        r.font.size = Pt(9)
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def run_generate(dry_run: bool = False, limit: int = 0) -> None:
    dsn = os.getenv("DATABASE_URL")
    if not dsn:
        raise RuntimeError("DATABASE_URL belum diset")

    with psycopg.connect(dsn) as conn:
        with conn.cursor() as cur:
            # Get surat with disposisi
            query = """
                SELECT
                    s.id,
                    s.unique_id,
                    s.agenda_ula,
                    s.surat_dari,
                    s.nomor_surat,
                    s.tgl_surat,
                    s.tgl_diterima_ula,
                    s.perihal,
                    s.arahan_menteri,
                    s.arahan_sekjen,
                    s.status_mailmerge,
                    l.dari_disposisi,
                    l.perihal_disposisi,
                    l.nomor_disposisi,
                    l.tanggal_disposisi
                FROM surat_dari_luar_bangda s
                LEFT JOIN lembar_disposisi_bangda l 
                    ON l.nomor_disposisi = s.agenda_ula
                WHERE s.agenda_ula IS NOT NULL
                ORDER BY s.id
            """
            if limit > 0:
                query += f" LIMIT {limit}"
            
            cur.execute(query)
            rows = cur.fetchall()
            cols = [d.name for d in cur.description]

        if not rows:
            log.info("Tidak ada data untuk diproses.")
            return

        if dry_run:
            log.info("[DRY-RUN] %d dokumen akan digenerate:", len(rows))
            for r in rows[:10]:
                row = dict(zip(cols, r))
                log.info("  %s | %s | %s", 
                        row.get("agenda_ula", "-"),
                        row.get("nomor_surat", "-"),
                        row.get("surat_dari", "-"))
            if len(rows) > 10:
                log.info("  ... dan %d lainnya", len(rows) - 10)
            return

        processed = failed = 0

        for r in rows:
            row = dict(zip(cols, r))
            safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', 
                             f"{row.get('agenda_ula', 'unknown')}_{row.get('id', 'unknown')}")
            file_name = f"Disposisi_ULA_{safe_name}.docx"
            local_path = os.path.join(LOCAL_DOCS_DIR, file_name)

            # Build disposisi content
            disposisi_content = ""
            if row.get("nomor_disposisi"):
                disposisi_content += f"No. Disposisi: {row['nomor_disposisi']}\n"
            if row.get("dari_disposisi"):
                disposisi_content += f"Dari: {row['dari_disposisi']}\n"
            if row.get("perihal_disposisi"):
                disposisi_content += f"Perihal: {row['perihal_disposisi']}"

            data = {
                "agenda_ula": row.get("agenda_ula"),
                "surat_dari": row.get("surat_dari"),
                "nomor_surat": row.get("nomor_surat"),
                "tgl_surat": row.get("tgl_surat"),
                "tgl_diterima_ula": row.get("tgl_diterima_ula"),
                "perihal": row.get("perihal"),
                "arahan_menteri": row.get("arahan_menteri"),
                "arahan_sekjen": row.get("arahan_sekjen"),
                "status_mailmerge": row.get("status_mailmerge"),
                "disposisi_content": disposisi_content
            }

            try:
                docx_bytes = create_disposisi_document(data)
                with open(local_path, "wb") as f:
                    f.write(docx_bytes)
                
                # Update DB dengan status
                with conn.cursor() as cur2:
                    cur2.execute("""
                        UPDATE surat_dari_luar_bangda
                        SET status_mailmerge = 'generated',
                            updated_at = NOW()
                        WHERE id = %s
                    """, (row['id'],))
                conn.commit()

                log.info("OK: %s → %s", row.get("agenda_ula", "-"), file_name)
                processed += 1

            except Exception as e:
                log.error("GAGAL %s: %s", row.get("agenda_ula", "-"), e)
                failed += 1

    result = {
        "ok": True,
        "processed": processed,
        "failed": failed,
        "output_dir": LOCAL_DOCS_DIR,
        "finished_at": datetime.now().isoformat()
    }
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--limit", type=int, default=0, help="Limit jumlah dokumen")
    args = parser.parse_args()
    run_generate(dry_run=args.dry_run, limit=args.limit)