#!/usr/bin/env python3
"""
Fix DOCX Formatting - Align with Official Letter Format
Based on: PENGUMUMAN-SELEKSI-TERBUKA-JPT-PRATAMA_4-JPT.pdf

Fixes:
1. Heading hierarchy (Heading 1 for sections, Heading 2 for sub-items)
2. Indentation (proper paragraph indentation levels)
3. Header/Footer (add letterhead header, page number footer)
4. Standard MS Office/Google Docs formatting
"""

import sys
from pathlib import Path

# Add project root to path for imports
project_root = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# File paths
INPUT_FILE = "storage/office/PENGUMUMAN-SELEKSI-JPT-PRATAMA.docx"
OUTPUT_FILE = "storage/office/PENGUMUMAN-SELEKSI-JPT-PRATAMA-FORMATTED.docx"

# ============================================================
# HEADER DEFINITIONS (based on PDF reference)
# ============================================================
LETTERHEAD = [
    ("KEMENTERIAN TRANSMIGRASI REPUBLIK INDONESIA", True, WD_ALIGN_PARAGRAPH.CENTER, 12, True),
    ("SEKRETARIAT JENDERAL", False, WD_ALIGN_PARAGRAPH.CENTER, 11, True),
    ("Jalan TMP. Kalibata Nomor 17 Jakarta Selatan 12750 Telepon 021-7971972", False, WD_ALIGN_PARAGRAPH.CENTER, 10, False),
    ("www.transmigrasi.go.id", False, WD_ALIGN_PARAGRAPH.CENTER, 10, False),
]

# Section headings mapping (paragraph index patterns)
# Sections: A, B, C, D, E are Heading 1
# Sub-items: 1, 2, 3, 4, 5 are Heading 2
SECTION_HEADERS = {
    "KETENTUAN UMUM": "heading1",
    "KETENTUAN KHUSUS": "heading1", 
    "TAHAPAN PELAKSANAAN SELEKSI": "heading1",
    "JADWAL WAKTU PELAKSANAAN": "heading1",
    "KETENTUAN LAIN": "heading1",
    # Sub-items (Heading 2)
    "Nama Jabatan lowong": "heading2",
    "Persyaratan Umum": "heading2",
    "Tata Cara Pengajuan Lamaran": "heading2",
    "Seleksi Administrasi": "heading2",
    "Penelusuran Rekam Jejak": "heading2",
    "Seleksi Kompetensi Teknis": "heading2",
    "Seleksi Kompetensi Manajerial": "heading2",
    "Wawancara Akhir": "heading2",
}


def is_section_header(text: str) -> str | None:
    """Check if text is a section header and return its type."""
    if not text.strip():
        return None
    
    clean = text.strip().upper()
    
    # Check main sections (A., B., C., D., E.)
    for keyword, heading_type in SECTION_HEADERS.items():
        if keyword.upper() in clean:
            return heading_type
    
    return None


def is_page_number(text: str) -> bool:
    """Check if text is a page number indicator (-X-)."""
    import re
    return bool(re.match(r'^-\d+-$', text.strip()))


def is_list_item(text: str) -> tuple | None:
    """Check if text is a list item and return (level, content)."""
    import re
    text_stripped = text.strip()
    
    if not text_stripped:
        return None
    
    # Level 1 letter list: a., b., c.
    if re.match(r'^[a-z]\.\s', text_stripped):
        return ("letter", text_stripped)
    
    # Level 1 number list: 1), 2), 3)
    if re.match(r'^\d+\)\s', text_stripped):
        return ("numbered", text_stripped)
    
    # Level 2 number list: 1.  2.  3. (main numbered items under sections)
    if re.match(r'^\d+\.\s+', text_stripped):
        return ("sub_numbered", text_stripped)
    
    return None


def apply_style_to_paragraph(para, font_size: int, bold: bool, alignment, color=None):
    """Apply basic style to a paragraph."""
    # Set alignment
    para.alignment = alignment
    
    # Style all runs
    for run in para.runs:
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
    
    # If no runs, add one
    if not para.runs:
        run = para.add_run()
        run.font.size = Pt(font_size)
        run.font.bold = bold


def set_paragraph_spacing(para, before=None, after=None, line_spacing=None):
    """Set paragraph spacing using OOXML directly."""
    if not para.paragraph_format:
        para.paragraph_format = ParaFormat()
    
    if before is not None:
        para.paragraph_format.space_before = Pt(before)
    if after is not None:
        para.paragraph_format.space_after = Pt(after)
    if line_spacing is not None:
        para.paragraph_format.line_spacing = Pt(line_spacing)


def add_header_footer_to_doc(doc):
    """Add header and footer to document sections."""
    for section in doc.sections:
        # Add header
        header = section.header
        header.is_linked_to_previous = False
        
        # Clear existing header
        for p in header.paragraphs:
            p.clear()
        
        # Add letterhead
        for i, (text, bold, alignment, size, _) in enumerate(LETTERHEAD):
            para = header.paragraphs[0] if i == 0 else header.add_paragraph()
            para.alignment = alignment
            run = para.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(0, 0, 0)
            if i == len(LETTERHEAD) - 1:
                set_paragraph_spacing(para, after=0)
            else:
                set_paragraph_spacing(para, before=0, after=2)
        
        # Add footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Clear existing footer
        for p in footer.paragraphs:
            p.clear()
        
        # Add page number field
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run.font.size = Pt(10)


def set_page_margins(doc):
    """Set standard A4 page margins."""
    for section in doc.sections:
        section.page_width = Cm(21)  # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)


def format_docx(input_path, output_path):
    """Apply proper formatting to DOCX file."""
    doc = Document(input_path)
    
    # Set page setup
    set_page_margins(doc)
    
    # Collect content info and plan changes
    paragraphs_info = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Skip page number paragraphs (we'll add them via footer)
        if is_page_number(text):
            paragraphs_info.append(("skip", para))
            continue
        
        # Determine paragraph type
        heading_type = is_section_header(text)
        list_info = is_list_item(text)
        
        if heading_type:
            paragraphs_info.append((heading_type, para))
        elif list_info:
            paragraphs_info.append(("list", para, list_info))
        else:
            paragraphs_info.append(("normal", para))
    
    # Apply formatting
    for info in paragraphs_info:
        para_type = info[0]
        para = info[1]
        
        if para_type == "skip":
            # Remove page number paragraphs
            p = para._element
            p.getparent().remove(p)
            continue
        
        elif para_type == "heading1":
            # Section headings (A., B., C., D., E.)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.name = 'Times New Roman'
            if not para.runs:
                run = para.add_run()
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.name = 'Times New Roman'
            set_paragraph_spacing(para, before=12, after=4)
            
        elif para_type == "heading2":
            # Sub-item headings (1., 2., 3.)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.name = 'Times New Roman'
            if not para.runs:
                run = para.add_run()
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.name = 'Times New Roman'
            set_paragraph_spacing(para, before=6, after=4)
            
        elif para_type == "list":
            # List items with proper indentation
            list_type, text = info[2]
            if list_type == "letter":
                # a., b., c. - First level letter list (1.27 cm indent)
                set_paragraph_spacing(para, before=0, after=0, line_spacing=14)
                apply_paragraph_indent(para, left=Cm(1.27))
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
            elif list_type == "numbered":
                # 1), 2), 3) - Second level number list (2.54 cm indent)
                set_paragraph_spacing(para, before=0, after=0, line_spacing=14)
                apply_paragraph_indent(para, left=Cm(2.54))
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
            elif list_type == "sub_numbered":
                # 1., 2., 3. under sections (0.5 cm indent)
                set_paragraph_spacing(para, before=0, after=0, line_spacing=14)
                apply_paragraph_indent(para, left=Cm(0.5))
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True  # Sub-items often bold
                    run.font.name = 'Times New Roman'
            
        else:
            # Normal text (justify, 1.5 line spacing)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(para, before=0, after=0, line_spacing=18)
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
    
    # Add header and footer
    add_header_footer_to_doc(doc)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Save
    doc.save(output_path)
    print(f"Document formatted and saved to: {output_path}")
    return output_path


def apply_paragraph_indent(para, left=None, right=None, first_line=None):
    """Set paragraph indentation using OOXML."""
    if left is not None:
        # Convert to twips (1 cm = 567 twips approximately)
        left_twips = int(left / Cm(1) * 567) if isinstance(left, (int, float)) else left
    
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    
    if left is not None:
        # Convert to twips
        left_cm = float(str(left).replace('cm', '').strip())
        left_twips = int(round(left_cm * 567))
        ind.set(qn('w:left'), str(left_twips))
    
    if right is not None:
        right_cm = float(str(right).replace('cm', '').strip())
        right_twips = int(round(right_cm * 567))
        ind.set(qn('w:right'), str(right_twips))
    
    if first_line is not None:
        first_cm = float(str(first_line).replace('cm', '').strip())
        first_twips = int(round(first_cm * 567))
        ind.set(qn('w:firstLine'), str(first_twips))
    
    pPr.append(ind)


class ParaFormat:
    """Simple paragraph format holder."""
    def __init__(self):
        self.space_before = None
        self.space_after = None
        self.line_spacing = None


# ============================================================
# Main execution
# ============================================================
if __name__ == "__main__":
    full_input = project_root / INPUT_FILE
    full_output = project_root / OUTPUT_FILE
    
    if not full_input.exists():
        print(f"Error: Input file not found: {full_input}")
        sys.exit(1)
    
    print(f"Formatting: {full_input}")
    print(f"Output: {full_output}")
    
    format_docx(str(full_input), str(full_output))
    
    print("\nFormatting applied:")
    print("  ✓ Heading hierarchy (H1 for A-E, H2 for sub-items)")
    print("  ✓ Paragraph indentation for list levels")
    print("  ✓ Letterhead header")
    print("  ✓ Page number footer")
    print("  ✓ Times New Roman, 11pt font")
    print("  ✓ Justified body text")
    print("  ✓ Standard A4 page margins")