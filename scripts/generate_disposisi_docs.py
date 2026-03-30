#!/usr/bin/env python3
"""
Generator Dokumen Disposisi — Pola DB-centric
===============================================
Alur baru:
  1. Export template dari Google Drive sebagai DOCX (read-only, tidak pakai quota)
  2. Isi placeholder menggunakan python-docx (proses lokal)
  3. Simpan DOCX ke /data/storage/disposisi_docs/
  4. Update DB (disposisi_documents) dengan local_file_path + generation_status='local_ready'
  5. Sync ke Drive dilakukan terpisah via sync_disposisi_to_gdrive.py

Jalankan:
  python3 scripts/generate_disposisi_docs.py
  python3 scripts/generate_disposisi_docs.py --dry-run
"""
import io
import os
import sys
import json
import argparse
import logging
import re
from datetime import datetime, date
from typing import Any, Dict, Optional
from copy import deepcopy

import psycopg
from docx import Document

PROJECT_ROOT = "/home/aseps/MCP/mcp-unified"
sys.path.insert(0, PROJECT_ROOT)
from core.secrets import load_runtime_secrets
load_runtime_secrets()

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("generate_disposisi")

# --- Mapping kode DARI → nama lengkap ---
_DARI_RAW = {
    "BU": "Bagian Umum",
    "UM": "Bagian Umum",
    "TU": "Tata Usaha",
    "PRC": "Bagian Perencanaan",
    "PUU": "Substansi Perundang-Undangan",
    "KEU": "Bagian Keuangan",
    "SD I": "Subdit Wilayah I",
    "SD II": "Subdit Wilayah II",
    "SD III": "Subdit Wilayah III",
    "SD IV": "Subdit Wilayah IV",
    "SD V": "Subdit Wilayah V",
    "SD VI": "Subdit Wilayah VI",
    "SD PMIPD": "Subdit PMIPD",
    "SD PIMPD": "Subdit PMIPD",
    "SD": "Subdit",
    "PEIPD": "Direktorat PEIPD",
    "PMIPD": "Subdit PMIPD",
    "SUPD I": "Direktorat SUPD I",
    "SUPD II": "Direktorat SUPD II",
    "SUPD III": "Direktorat SUPD III",
    "SUPD IV": "Direktorat SUPD IV",
    "PPK": "Pejabat Pembuat Komitmen",
    "BANGDA": "Ditjen Bina Pembangunan Daerah",
    "TU SUPD II": "Tata Usaha SUPD II",
}
# Normalisasi key: uppercase, dots/spaces → single space
import re as _re
DARI_LOOKUP: dict = {}
for _k, _v in _DARI_RAW.items():
    _norm = _re.sub(r'[\s\.]+', ' ', _k.upper()).strip()
    DARI_LOOKUP[_norm] = _v


def map_dari(dari: str) -> str:
    """Kembalikan nama lengkap dari kode DARI, fallback ke nilai asli."""
    if not dari:
        return ""
    norm = re.sub(r'[\s\.]+', ' ', dari.strip().upper()).strip()
    return DARI_LOOKUP.get(norm, dari.strip())


# --- Konfigurasi ---
TEMPLATE_ID     = "1ixgD-8ISGkyaD018sNfEPznCFvxCOFC5xbGE5GfQVoQ"
FOLDER_ID       = "1s1WyweDstV0vYgP1SIfQk4rWwDGO0OYw"
LOCAL_DOCS_DIR  = "/home/aseps/MCP/data/storage/disposisi_docs"
SA_KEY_FILE     = "/home/aseps/MCP/config/credentials/google/mcp-gmail-482015-682b788ee191.json"
SA_SCOPES       = ["https://www.googleapis.com/auth/drive.readonly"]

os.makedirs(LOCAL_DOCS_DIR, exist_ok=True)


# --- Helper: format tanggal ---
def fmt_date(d: Any) -> str:
    BULAN = {1:"Januari",2:"Februari",3:"Maret",4:"April",5:"Mei",6:"Juni",
             7:"Juli",8:"Agustus",9:"September",10:"Oktober",11:"November",12:"Desember"}
    if isinstance(d, (datetime, date)):
        return f"{d.day:02d} {BULAN.get(d.month,'')} {d.year}"
    return str(d) if d else ""


# --- Export template DOCX dari Drive (sekali, dicache) ---
_template_bytes: Optional[bytes] = None

def get_template_bytes() -> bytes:
    global _template_bytes
    if _template_bytes is not None:
        return _template_bytes
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    creds = service_account.Credentials.from_service_account_file(SA_KEY_FILE, scopes=SA_SCOPES)
    drive = build("drive", "v3", credentials=creds, cache_discovery=False)
    log.info("Mengunduh template DOCX dari Drive...")
    _template_bytes = drive.files().export(
        fileId=TEMPLATE_ID,
        mimeType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ).execute()
    log.info("Template DOCX: %d bytes", len(_template_bytes))
    return _template_bytes


# --- Ganti placeholder dalam DOCX ---
def replace_in_runs(paragraph, replacements: Dict[str, str]) -> None:
    """Ganti <<KEY>> dalam teks paragraph, aman meski key terpecah antar runs."""
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    for key, val in replacements.items():
        new_text = new_text.replace(f"<<{key}>>", val or "")
    if new_text == full_text:
        return
    # Tulis ke run pertama, kosongkan sisanya
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def fill_docx(template_bytes: bytes, replacements: Dict[str, str]) -> bytes:
    """Isi placeholder dalam DOCX template, return bytes hasil."""
    doc = Document(io.BytesIO(template_bytes))

    # Proses semua paragraf (termasuk di dalam tabel)
    def process_paragraphs(paragraphs):
        for para in paragraphs:
            replace_in_runs(para, replacements)

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- Generate semua dokumen ---
def run_generate(dry_run: bool = False) -> None:
    dsn = os.getenv("DATABASE_URL")
    if not dsn:
        raise RuntimeError("DATABASE_URL belum diset")

    with psycopg.connect(dsn) as conn:
        with conn.cursor() as cur:
            cur.execute("""
                SELECT
                    ld.id           AS ld_id,
                    ld.agenda_puu,
                    ld.direktorat,
                    ld.tgl_diterima,
                    sp.nomor_nd,
                    sp.hal,
                    sp.tanggal_surat,
                    sp.no_agenda_dispo,
                    sp.dari
                FROM lembar_disposisi ld
                JOIN surat_masuk_puu sp ON sp.id = ld.surat_id
                LEFT JOIN disposisi_documents dd ON dd.lembar_disposisi_id = ld.id
                WHERE ld.agenda_puu IS NOT NULL
                  AND ld.agenda_puu NOT LIKE '%http%'
                  AND (dd.id IS NULL
                       OR dd.generation_status IN ('failed','pending')
                       OR dd.local_file_path IS NULL)
                ORDER BY ld.id
            """)
            rows = cur.fetchall()
            cols = [d.name for d in cur.description]

        if not rows:
            log.info("Semua dokumen sudah digenerate.")
            print(json.dumps({"ok": True, "processed": 0, "message": "Tidak ada yang perlu diproses"}))
            return

        if dry_run:
            log.info("[DRY-RUN] %d dokumen akan digenerate:", len(rows))
            for r in rows:
                row = dict(zip(cols, r))
                log.info("  %s | %s", row["agenda_puu"], row["nomor_nd"])
            return

        template_bytes = get_template_bytes()
        processed = failed = 0

        for r in rows:
            row = dict(zip(cols, r))
            agenda_puu  = row["agenda_puu"]
            ld_id       = row["ld_id"]
            # Nama file aman untuk filesystem
            safe_name   = re.sub(r'[\\/:*?"<>|]', '_', agenda_puu)
            file_name   = f"Disposisi - {safe_name}.docx"
            local_path  = os.path.join(LOCAL_DOCS_DIR, file_name)

            # Format "Surat Dari": "Bagian Umum - Sekretariat"
            dari_full   = map_dari(row["dari"] or "")
            direktorat  = row["direktorat"] or ""
            # Hindari duplikasi: "Direktorat PEIPD - Direktorat PEIPD"
            surat_dari  = (
                f"{dari_full} - {direktorat}"
                if dari_full and direktorat and dari_full.lower() != direktorat.lower()
                else (dari_full or direktorat)
            )

            replacements = {
                "DIREKTORAT":    surat_dari,
                "NOMOR ND":      row["nomor_nd"] or "",
                "TANGGAL_SURAT": fmt_date(row["tanggal_surat"]),
                "HAL":           row["hal"] or "",
                "TGL DITERIMA":  fmt_date(row["tgl_diterima"]),
                "NOMOR_ND":      row["no_agenda_dispo"] or "",
                "AGENDA_PUU":    agenda_puu,
            }

            try:
                # Generate DOCX lokal
                docx_bytes = fill_docx(template_bytes, replacements)
                with open(local_path, "wb") as f:
                    f.write(docx_bytes)

                # Update DB
                with conn.cursor() as cur:
                    cur.execute("""
                        INSERT INTO disposisi_documents
                            (lembar_disposisi_id, agenda_puu, file_name,
                             folder_id, template_id, local_file_path,
                             generated_at, generation_status, sync_status)
                        VALUES (%s,%s,%s,%s,%s,%s,NOW(),'local_ready','pending')
                        ON CONFLICT (lembar_disposisi_id) DO UPDATE SET
                            local_file_path   = EXCLUDED.local_file_path,
                            generated_at      = NOW(),
                            generation_status = 'local_ready',
                            sync_status       = 'pending'
                    """, (ld_id, agenda_puu, file_name, FOLDER_ID, TEMPLATE_ID, local_path))
                conn.commit()

                log.info("OK: %s → %s", agenda_puu, local_path)
                processed += 1

            except Exception as e:
                log.error("GAGAL %s: %s", agenda_puu, e)
                with conn.cursor() as cur:
                    cur.execute("""
                        INSERT INTO disposisi_documents
                            (lembar_disposisi_id, agenda_puu, file_name,
                             folder_id, template_id, generation_status,
                             sync_status, error_message)
                        VALUES (%s,%s,%s,%s,%s,'failed','error',%s)
                        ON CONFLICT (lembar_disposisi_id) DO UPDATE SET
                            generation_status = 'failed',
                            error_message     = EXCLUDED.error_message
                    """, (ld_id, agenda_puu, file_name, FOLDER_ID, TEMPLATE_ID, str(e)))
                conn.commit()
                failed += 1

    print(json.dumps({
        "ok": True,
        "processed": processed,
        "failed": failed,
        "output_dir": LOCAL_DOCS_DIR,
        "finished_at": datetime.utcnow().isoformat()
    }, indent=2))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()
    run_generate(dry_run=args.dry_run)
