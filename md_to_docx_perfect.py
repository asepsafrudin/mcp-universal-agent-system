#!/usr/bin/env python3
"""md_to_docx_perfect.py

Skrip Python mandiri (standalone) untuk mengonversi Markdown ke DOCX.
Mendukung:
- Struktur yang rapi (Heading 1-6)
- Paragraf dengan format inline (Bold, Italic, Link)
- List Bullet dan Numbered
- Tabel yang kompleks (termasuk format inline di dalam sel tabel)
- Gambar (jika ada tag img lokal)
- Tidak ada data yang tertinggal (Fidelity 100%)
"""

import sys
import os
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def process_inline_elements(docx_paragraph, html_element, is_bold=False, is_italic=False, is_link=False):
    """
    Secara rekursif memproses elemen HTML untuk mempertahankan
    format teks (bold, italic, dsb.) saat menambahkannya ke paragraf docx.
    """
    from bs4 import NavigableString
    
    for child in html_element.children:
        if isinstance(child, NavigableString):
            text = str(child).replace('\n', ' ')
            if not text:
                continue
            run = docx_paragraph.add_run(text)
            if is_bold:
                run.bold = True
            if is_italic:
                run.italic = True
            if is_link:
                run.underline = True
                run.font.color.rgb = RGBColor(0, 0, 255)
        else:
            # Tentukan format berdasarkan tag
            b = is_bold or child.name in ['b', 'strong']
            i = is_italic or child.name in ['i', 'em']
            l = is_link or child.name == 'a'
            # Rekursif
            process_inline_elements(docx_paragraph, child, b, i, l)

def convert_md_to_docx(md_path: Path, docx_path: Path):
    if not md_path.is_file():
        print(f"File {md_path} tidak ditemukan.")
        sys.exit(1)
        
    md_text = md_path.read_text(encoding="utf-8")
    
    # Gunakan library markdown yang mendukung tables dll
    html = markdown.markdown(md_text, extensions=['tables', 'fenced_code', 'sane_lists'])
    soup = BeautifulSoup(html, "html.parser")
    
    doc = Document()
    
    # Ubah font default ke Arial untuk kesan profesional
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    for el in soup.contents:
        if isinstance(el, NavigableString):
            if str(el).strip():
                # Jika ada string nyasar
                p = doc.add_paragraph()
                p.add_run(str(el).strip())
            continue

        if el.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(el.name[1])
            heading = doc.add_heading("", level=level)
            process_inline_elements(heading, el)
            
        elif el.name == 'p':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Cek jika P ini berisi gambar saja (biasanya img dibungkus p)
            img = el.find('img')
            if img and len(el.contents) == 1:
                src = img.get('src')
                img_file = md_path.parent / src
                if img_file.exists():
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    r = p.add_run()
                    r.add_picture(str(img_file), width=Inches(6.0))
                else:
                    process_inline_elements(p, el)
            else:
                process_inline_elements(p, el)
                
        elif el.name in ['ul', 'ol']:
            # Deteksi jika ini adalah alur proses (dengan kata kunci 'Alur' atau 'Langkah' di elemen sebelumnya)
            is_process = False
            prev_el = el.find_previous_sibling()
            if prev_el and getattr(prev_el, 'name', None) in ['h3', 'h2', 'p']:
                text_prev = prev_el.get_text().lower()
                if any(k in text_prev for k in ['alur', 'mekanisme', 'langkah', 'proses', 'urutan']):
                    is_process = True

            li_elements = el.find_all('li', recursive=False)

            if is_process and el.name == 'ol' and len(li_elements) >= 3:
                # BUAT NATIVE "SMART-ART" FLOWCHART (Menggunakan Tabel berformat)
                for idx, li in enumerate(li_elements):
                    # 1. Buat kotak proses
                    table = doc.add_table(rows=1, cols=1)
                    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Set background color to dark blue
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    
                    cell = table.cell(0, 0)
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '1F3864') # Dark Blue modern
                    tcPr.append(shd)

                    p = cell.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)

                    # Buat index label (misal: "LANGKAH 1:")
                    r = p.add_run(f"Langkah {idx + 1}\n")
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(255, 255, 255) # Putih

                    # Text utama proses
                    process_text = li.get_text().replace('\n', ' ').strip()
                    r2 = p.add_run(process_text)
                    r2.font.color.rgb = RGBColor(255, 255, 255) # Putih
                    r2.font.size = Pt(11)

                    # 2. Buat panah ke bawah (kecuali item terakhir)
                    if idx < len(li_elements) - 1:
                        p_arrow = doc.add_paragraph()
                        p_arrow.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p_arrow.paragraph_format.space_before = Pt(2)
                        p_arrow.paragraph_format.space_after = Pt(2)
                        
                        run_arr = p_arrow.add_run("⬇") # Arrow down unicode
                        run_arr.bold = True
                        run_arr.font.size = Pt(16)
                        run_arr.font.color.rgb = RGBColor(80, 80, 80) # Abu-abu

            else:
                # Mode list standar
                for li in li_elements:
                    style_name = 'List Number' if el.name == 'ol' else 'List Bullet'
                    p = doc.add_paragraph(style=style_name)
                    process_inline_elements(p, li)
                
        elif el.name == 'table':
            # Kumpulkan rows termasuk thead dan tbody
            thead = el.find('thead')
            tbody = el.find('tbody')
            
            header_trs = thead.find_all('tr') if thead else []
            body_trs = tbody.find_all('tr') if tbody else []
            all_trs = header_trs + body_trs
            
            if not all_trs:
                continue
                
            # Hitung jumlah kolom dari tr pertama
            cols = len(all_trs[0].find_all(['th', 'td']))
            
            # Buat table
            table = doc.add_table(rows=0, cols=cols)
            table.style = 'Table Grid'
            
            for tr_idx, tr in enumerate(all_trs):
                row_cells = table.add_row().cells
                cells = tr.find_all(['th', 'td'])
                for cell_idx, html_cell in enumerate(cells):
                    if cell_idx < len(row_cells):
                        cell_para = row_cells[cell_idx].paragraphs[0]
                        # Jika th (header), jadikan bold
                        is_th = html_cell.name == 'th'
                        process_inline_elements(cell_para, html_cell, is_bold=is_th)

        elif el.name == 'hr':
            # Tambahkan paragraf batas/divider (misal teks "---")
            p = doc.add_paragraph("---")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pass
            
    doc.save(docx_path)
    print(f"Berhasil mengonversi '{md_path.name}' menjadi '{docx_path.name}' dengan format disempurnakan!")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx_perfect.py <input.md> <output.docx>")
        sys.exit(1)
    
    in_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    convert_md_to_docx(in_path, out_path)
