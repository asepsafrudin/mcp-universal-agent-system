"""
DOCX Extractor

Mengekstrak teks dari Microsoft Word documents (.docx, .doc).
"""

from pathlib import Path
from typing import Dict, Any, List


class DocxExtractor:
    """
    Extractor untuk file Microsoft Word.
    
    Supports:
        - .docx (Word 2007+)
        - .doc (legacy, via conversion)
    """
    
    def __init__(self):
        """Initialize DOCX extractor."""
        self._lib = None
    
    def _init_lib(self):
        """Lazy initialization library."""
        if self._lib is None:
            try:
                import docx
                self._lib = "python-docx"
            except ImportError:
                raise ImportError(
                    "python-docx tidak tersedia. "
                    "Install: pip install python-docx"
                )
    
    async def extract(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text dari DOCX file.
        
        Args:
            file_path: Path ke .docx file
            
        Returns:
            Dict dengan keys:
                - text: Extracted text
                - metadata: Document metadata
                - tables: List of tables (if any)
        """
        self._init_lib()
        
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File tidak ditemukan: {file_path}")
        
        if path.suffix.lower() == ".doc":
            # Legacy .doc format - perlu conversion
            return await self._extract_legacy_doc(file_path)
        
        return await self._extract_docx(file_path)
    
    async def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract dari .docx file."""
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract metadata
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "file_type": "docx"
        }
        
        # Extract text dari paragraphs
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if heading
                if para.style.name.startswith('Heading'):
                    text_parts.append(f"\n## {para.text}\n")
                else:
                    text_parts.append(para.text)
        
        # Extract tables
        tables_data = []
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            # Convert table ke text format
            if table_data:
                text_parts.append(f"\n--- Table {table_idx + 1} ---")
                for row in table_data:
                    text_parts.append(" | ".join(row))
                text_parts.append("")
            
            tables_data.append(table_data)
        
        full_text = "\n".join(text_parts)
        
        return {
            "text": full_text,
            "metadata": metadata,
            "tables": tables_data,
            "ocr_used": False
        }
    
    async def _extract_legacy_doc(self, file_path: str) -> Dict[str, Any]:
        """
        Extract dari legacy .doc file.
        
        Note: Memerlukan conversion atau antiword.
        Ini adalah placeholder.
        """
        # TODO: Implementasi menggunakan antiword atau conversion
        # For now, return error dengan info
        return {
            "text": "",
            "metadata": {
                "file_type": "doc",
                "error": "Legacy .doc format memerlukan conversion. "
                        "Gunakan .docx atau convert dulu menggunakan antiword."
            },
            "tables": [],
            "ocr_used": False
        }