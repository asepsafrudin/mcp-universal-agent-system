"""
DOCX Tools for reading and writing Microsoft Word documents
"""
import docx
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tools.base import BaseTool, register_tool


@register_tool
def read_docx(file_path: str) -> Dict:
    """
    Read a DOCX file and return its content structure
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        Dictionary with paragraphs, tables, and metadata
    """
    try:
        doc = Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'alignment': str(para.alignment) if para.alignment else 'LEFT'
                })
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        # Extract metadata
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'core_properties': {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': str(doc.core_properties.created) if doc.core_properties.created else None,
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else None,
            }
        }
        
        return {
            'success': True,
            'file_path': file_path,
            'paragraphs': paragraphs,
            'tables': tables,
            'metadata': metadata
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def write_docx(
    file_path: str,
    content: List[Dict],
    title: Optional[str] = None,
    author: Optional[str] = None
) -> Dict:
    """
    Create a new DOCX file with specified content
    
    Args:
        file_path: Path where the .docx file will be saved
        content: List of content items (paragraphs, tables)
               Each item should have 'type' ('paragraph' or 'table') and 'data'
        title: Document title (optional)
        author: Document author (optional)
        
    Returns:
        Success status and file path
    """
    try:
        doc = Document()
        
        # Set metadata
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Add content
        for item in content:
            item_type = item.get('type', 'paragraph')
            
            if item_type == 'paragraph':
                text = item.get('text', '')
                style = item.get('style', 'Normal')
                alignment = item.get('alignment', 'LEFT')
                
                para = doc.add_paragraph(text, style=style)
                
                # Set alignment
                if alignment == 'CENTER':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignment == 'RIGHT':
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif alignment == 'JUSTIFY':
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
            elif item_type == 'table':
                table_data = item.get('data', [])
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    for i, row_data in enumerate(table_data):
                        for j, cell_text in enumerate(row_data):
                            table.rows[i].cells[j].text = str(cell_text)
                            
            elif item_type == 'heading':
                text = item.get('text', '')
                level = item.get('level', 1)
                doc.add_heading(text, level=level)
        
        # Save document
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'message': f'Document saved successfully with {len(content)} items'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def extract_text_docx(file_path: str, include_tables: bool = True) -> Dict:
    """
    Extract all text content from a DOCX file
    
    Args:
        file_path: Path to the .docx file
        include_tables: Whether to include table content in extraction
        
    Returns:
        Dictionary with extracted text and statistics
    """
    try:
        doc = Document(file_path)
        
        # Extract paragraph text
        paragraphs_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_text.append(para.text)
        
        # Extract table text
        tables_text = []
        if include_tables:
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_rows.append(' | '.join(row_text))
                tables_text.append('\n'.join(table_rows))
        
        # Combine all text
        all_text = '\n\n'.join(paragraphs_text)
        if include_tables and tables_text:
            all_text += '\n\n[Tables]\n\n' + '\n\n'.join(tables_text)
        
        return {
            'success': True,
            'file_path': file_path,
            'text': all_text,
            'paragraph_count': len(paragraphs_text),
            'table_count': len(doc.tables),
            'total_characters': len(all_text),
            'total_words': len(all_text.split())
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def edit_docx(file_path: str, edits: List[Dict]) -> Dict:
    """
    Edit an existing DOCX file
    
    Args:
        file_path: Path to the .docx file
        edits: List of edit operations
               Each edit should have:
               - 'action': 'replace', 'insert', 'delete', 'append'
               - 'target': target paragraph index or 'end'
               - 'content': new content (for replace/insert/append)
               
    Returns:
        Success status and details of changes
    """
    try:
        doc = Document(file_path)
        changes_made = []
        
        for edit in edits:
            action = edit.get('action')
            target = edit.get('target')
            content = edit.get('content', '')
            
            if action == 'append':
                doc.add_paragraph(content)
                changes_made.append(f"Appended paragraph: {content[:50]}...")
                
            elif action == 'insert' and isinstance(target, int):
                if 0 <= target < len(doc.paragraphs):
                    doc.paragraphs[target].insert_paragraph_before(content)
                    changes_made.append(f"Inserted at position {target}")
                    
            elif action == 'replace' and isinstance(target, int):
                if 0 <= target < len(doc.paragraphs):
                    doc.paragraphs[target].text = content
                    changes_made.append(f"Replaced paragraph {target}")
                    
            elif action == 'delete' and isinstance(target, int):
                if 0 <= target < len(doc.paragraphs):
                    p = doc.paragraphs[target]
                    p._element.getparent().remove(p._element)
                    changes_made.append(f"Deleted paragraph {target}")
        
        # Save changes
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'changes_made': changes_made,
            'total_changes': len(changes_made)
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def search_replace_docx(
    file_path: str,
    search_text: str,
    replace_text: str,
    case_sensitive: bool = False,
    use_regex: bool = False
) -> Dict:
    """
    Search and replace text throughout a DOCX document
    
    Args:
        file_path: Path to the .docx file
        search_text: Text to search for
        replace_text: Text to replace with
        case_sensitive: Whether search is case-sensitive (default: False)
        use_regex: Whether to use regex pattern matching (default: False)
        
    Returns:
        Success status and details of replacements made
    """
    try:
        import re
        doc = Document(file_path)
        
        replacements_made = 0
        paragraphs_modified = []
        
        # Compile regex pattern if needed
        if use_regex:
            flags = 0 if case_sensitive else re.IGNORECASE
            pattern = re.compile(search_text, flags)
        
        for para_idx, para in enumerate(doc.paragraphs):
            original_text = para.text
            
            if use_regex:
                new_text = pattern.sub(replace_text, original_text)
            else:
                if case_sensitive:
                    new_text = original_text.replace(search_text, replace_text)
                else:
                    new_text = original_text.replace(search_text.lower(), replace_text)
                    # Also try exact match if different
                    if new_text == original_text:
                        new_text = original_text.replace(search_text, replace_text)
            
            if new_text != original_text:
                para.text = new_text
                replacements_made += 1
                paragraphs_modified.append(para_idx)
        
        # Also search in table cells
        tables_modified = 0
        for table_idx, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    
                    if use_regex:
                        new_text = pattern.sub(replace_text, original_text)
                    else:
                        if case_sensitive:
                            new_text = original_text.replace(search_text, replace_text)
                        else:
                            new_text = original_text.replace(search_text.lower(), replace_text)
                            if new_text == original_text:
                                new_text = original_text.replace(search_text, replace_text)
                    
                    if new_text != original_text:
                        cell.text = new_text
                        tables_modified += 1
        
        # Save changes
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'replacements_made': replacements_made,
            'paragraphs_modified': paragraphs_modified,
            'tables_modified': tables_modified,
            'search_text': search_text,
            'replace_text': replace_text,
            'case_sensitive': case_sensitive,
            'use_regex': use_regex
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def apply_paragraph_style_docx(
    file_path: str,
    paragraph_idx: int,
    font_name: Optional[str] = None,
    font_size: Optional[int] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    color: Optional[str] = None,
    alignment: Optional[str] = None,
    line_spacing: Optional[float] = None,
    space_before: Optional[float] = None,
    space_after: Optional[float] = None
) -> Dict:
    """
    Apply comprehensive styling to a paragraph in a DOCX document
    
    Args:
        file_path: Path to DOCX file
        paragraph_idx: Index of paragraph to style (0-based)
        font_name: Font family name (e.g., 'Arial', 'Times New Roman')
        font_size: Font size in points (e.g., 12, 14)
        bold: True for bold text
        italic: True for italic text
        underline: True for underlined text
        color: Font color in hex format (e.g., 'FF0000' for red, '0000FF' for blue)
        alignment: Text alignment ('LEFT', 'CENTER', 'RIGHT', 'JUSTIFY')
        line_spacing: Line spacing multiplier (e.g., 1.0, 1.5, 2.0)
        space_before: Space before paragraph in points
        space_after: Space after paragraph in points
        
    Returns:
        Success status and details of styles applied
    """
    try:
        from docx.shared import RGBColor
        
        doc = Document(file_path)
        
        if paragraph_idx >= len(doc.paragraphs):
            return {
                'success': False,
                'error': f'Paragraph index {paragraph_idx} out of range. Document has {len(doc.paragraphs)} paragraphs.'
            }
        
        para = doc.paragraphs[paragraph_idx]
        styles_applied = {}
        
        # Apply font styling to all runs in paragraph
        if any([font_name, font_size, bold, italic, underline, color]):
            # If paragraph has no runs, add one
            if not para.runs:
                para.add_run()
            
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                    styles_applied['font_name'] = font_name
                
                if font_size:
                    run.font.size = Pt(font_size)
                    styles_applied['font_size'] = font_size
                
                if bold is not None:
                    run.font.bold = bold
                    styles_applied['bold'] = bold
                
                if italic is not None:
                    run.font.italic = italic
                    styles_applied['italic'] = italic
                
                if underline is not None:
                    run.font.underline = underline
                    styles_applied['underline'] = underline
                
                if color:
                    # Convert hex to RGB
                    color = color.lstrip('#')
                    if len(color) == 6:
                        r = int(color[0:2], 16)
                        g = int(color[2:4], 16)
                        b = int(color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                        styles_applied['color'] = f'#{color}'
        
        # Apply paragraph formatting
        if alignment:
            alignment_map = {
                'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if alignment.upper() in alignment_map:
                para.alignment = alignment_map[alignment.upper()]
                styles_applied['alignment'] = alignment.upper()
        
        if line_spacing is not None:
            para.paragraph_format.line_spacing = line_spacing
            styles_applied['line_spacing'] = line_spacing
        
        if space_before is not None:
            para.paragraph_format.space_before = Pt(space_before)
            styles_applied['space_before'] = space_before
        
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)
            styles_applied['space_after'] = space_after
        
        # Save changes
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'paragraph_idx': paragraph_idx,
            'styles_applied': styles_applied
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


# ==================== P1 FEATURES - DOCX ====================

@register_tool
def add_header_footer_docx(
    file_path: str,
    header_text: Optional[str] = None,
    footer_text: Optional[str] = None,
    alignment: str = 'CENTER'
) -> Dict:
    """
    Add or modify header and footer in a DOCX document
    
    Args:
        file_path: Path to DOCX file
        header_text: Text for header (None to skip)
        footer_text: Text for footer (None to skip)
        alignment: Text alignment ('LEFT', 'CENTER', 'RIGHT')
        
    Returns:
        Success status and details
    """
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document(file_path)
        
        alignment_map = {
            'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
            'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT
        }
        align = alignment_map.get(alignment.upper(), WD_ALIGN_PARAGRAPH.CENTER)
        
        result = {'header_added': False, 'footer_added': False}
        
        # Add header
        if header_text:
            header = doc.sections[0].header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.text = header_text
            header_para.alignment = align
            result['header_added'] = True
            result['header_text'] = header_text
        
        # Add footer
        if footer_text:
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = footer_text
            footer_para.alignment = align
            result['footer_added'] = True
            result['footer_text'] = footer_text
        
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            **result
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def insert_image_docx(
    file_path: str,
    image_path: str,
    paragraph_idx: Optional[int] = None,
    width: Optional[float] = None,
    height: Optional[float] = None
) -> Dict:
    """
    Insert an image into a DOCX document
    
    Args:
        file_path: Path to DOCX file
        image_path: Path to image file (jpg, png, etc.)
        paragraph_idx: Index to insert after (None = append at end)
        width: Image width in inches (optional)
        height: Image height in inches (optional)
        
    Returns:
        Success status and details
    """
    try:
        from docx.shared import Inches
        
        doc = Document(file_path)
        
        # Prepare image dimensions
        img_width = Inches(width) if width else None
        img_height = Inches(height) if height else None
        
        # Insert at specified position or append
        if paragraph_idx is not None and paragraph_idx < len(doc.paragraphs):
            para = doc.paragraphs[paragraph_idx]
            run = para.add_run()
            run.add_picture(image_path, width=img_width, height=img_height)
        else:
            # Append at end
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(image_path, width=img_width, height=img_height)
        
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'image_path': image_path,
            'paragraph_idx': paragraph_idx,
            'width': width,
            'height': height
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def add_hyperlink_docx(
    file_path: str,
    paragraph_idx: int,
    text: str,
    url: str,
    tooltip: Optional[str] = None
) -> Dict:
    """
    Add a hyperlink to a paragraph in a DOCX document
    
    Args:
        file_path: Path to DOCX file
        paragraph_idx: Index of paragraph to add link to
        text: Display text for the hyperlink
        url: URL for the hyperlink
        tooltip: Hover tooltip (optional)
        
    Returns:
        Success status and details
    """
    try:
        doc = Document(file_path)
        
        if paragraph_idx >= len(doc.paragraphs):
            return {
                'success': False,
                'error': f'Paragraph index {paragraph_idx} out of range'
            }
        
        para = doc.paragraphs[paragraph_idx]
        
        # Create hyperlink
        part = doc.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        
        # Create hyperlink element
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        
        # Create run element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        
        # Create run properties (blue, underlined)
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(docx.oxml.shared.qn('w:val'), '0000FF')
        rPr.append(color)
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
        new_run.append(rPr)
        
        # Add text
        t = docx.oxml.shared.OxmlElement('w:t')
        t.text = text
        new_run.append(t)
        
        hyperlink.append(new_run)
        para._p.append(hyperlink)
        
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'paragraph_idx': paragraph_idx,
            'text': text,
            'url': url
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def add_list_docx(
    file_path: str,
    items: List[str],
    list_type: str = 'bullet',
    start_at: int = 1
) -> Dict:
    """
    Add a bulleted or numbered list to a DOCX document
    
    Args:
        file_path: Path to DOCX file (will be created if doesn't exist)
        items: List of text items
        list_type: 'bullet' or 'numbered'
        start_at: Starting number for numbered lists (default: 1)
        
    Returns:
        Success status and details
    """
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # Check if file exists, create if not
        try:
            doc = Document(file_path)
        except:
            doc = Document()
        
        items_added = []
        
        for i, item in enumerate(items):
            if list_type == 'bullet':
                para = doc.add_paragraph(item, style='List Bullet')
            else:  # numbered
                para = doc.add_paragraph(item, style='List Number')
            items_added.append(item)
        
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'list_type': list_type,
            'items_added': len(items_added),
            'items': items_added
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def set_page_setup_docx(
    file_path: str,
    orientation: str = 'portrait',
    width: Optional[float] = None,
    height: Optional[float] = None,
    margin_top: float = 1.0,
    margin_bottom: float = 1.0,
    margin_left: float = 1.0,
    margin_right: float = 1.0
) -> Dict:
    """
    Set page setup options for a DOCX document
    
    Args:
        file_path: Path to DOCX file
        orientation: 'portrait' or 'landscape'
        width: Page width in inches (optional, default based on orientation)
        height: Page height in inches (optional, default based on orientation)
        margin_top: Top margin in inches
        margin_bottom: Bottom margin in inches
        margin_left: Left margin in inches
        margin_right: Right margin in inches
        
    Returns:
        Success status and details
    """
    try:
        from docx.shared import Inches
        from docx.enum.section import WD_ORIENT
        
        doc = Document(file_path)
        section = doc.sections[0]
        
        # Set orientation
        if orientation.lower() == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            # Default landscape dimensions
            if width is None:
                width = 11.0
            if height is None:
                height = 8.5
        else:  # portrait
            section.orientation = WD_ORIENT.PORTRAIT
            # Default portrait dimensions
            if width is None:
                width = 8.5
            if height is None:
                height = 11.0
        
        # Set page size
        section.page_width = Inches(width)
        section.page_height = Inches(height)
        
        # Set margins
        section.top_margin = Inches(margin_top)
        section.bottom_margin = Inches(margin_bottom)
        section.left_margin = Inches(margin_left)
        section.right_margin = Inches(margin_right)
        
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'orientation': orientation,
            'width': width,
            'height': height,
            'margins': {
                'top': margin_top,
                'bottom': margin_bottom,
                'left': margin_left,
                'right': margin_right
            }
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }


@register_tool
def add_toc_docx(
    file_path: str,
    levels: int = 3,
    title: str = 'Table of Contents'
) -> Dict:
    """
    Add a Table of Contents to a DOCX document
    Note: TOC will be auto-generated when document is opened in Word
    
    Args:
        file_path: Path to DOCX file
        levels: Number of heading levels to include (1-9)
        title: Title for the TOC
        
    Returns:
        Success status and details
    """
    try:
        doc = Document(file_path)
        
        # Add TOC title
        doc.add_paragraph(title, style='Heading 1')
        
        # Add TOC field
        para = doc.add_paragraph()
        run = para.add_run()
        
        # Create TOC field code
        fldChar1 = docx.oxml.shared.OxmlElement('w:fldChar')
        fldChar1.set(docx.oxml.shared.qn('w:fldCharType'), 'begin')
        
        instrText = docx.oxml.shared.OxmlElement('w:instrText')
        instrText.set(docx.oxml.shared.qn('xml:space'), 'preserve')
        instrText.text = f'TOC \\o "1-{levels}" \\h \\z \\u'
        
        fldChar2 = docx.oxml.shared.OxmlElement('w:fldChar')
        fldChar2.set(docx.oxml.shared.qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        doc.save(file_path)
        
        return {
            'success': True,
            'file_path': file_path,
            'levels': levels,
            'title': title,
            'note': 'TOC will be populated when opened in Microsoft Word'
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'file_path': file_path
        }
