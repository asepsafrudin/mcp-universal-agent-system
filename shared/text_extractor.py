"""
Text Extractor untuk berbagai format dokumen.

Usage:
    from shared.text_extractor import extract_text
    
    text = extract_text("file.pdf")
    text = extract_text("file.docx")
"""
import io
import re
from pathlib import Path
from typing import Optional


def extract_text(file_path: str | Path) -> str:
    """
    Extract text dari file berdasarkan extension.
    
    Args:
        file_path: Path ke file
        
    Returns:
        Extracted text atau error message
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    
    try:
        if ext == '.pdf':
            return extract_pdf_text(file_path)
        elif ext in ['.docx', '.doc']:
            return extract_docx_text(file_path)
        elif ext in ['.txt', '.md']:
            return extract_txt_text(file_path)
        else:
            return f"[Unsupported file type: {ext}]"
    except Exception as e:
        return f"[Extraction error: {e}]"


def extract_pdf_text(file_path: Path) -> str:
    """Extract text dari PDF menggunakan PyPDF2 atau pdfplumber."""
    try:
        # Coba pdfplumber (lebih baik untuk layout)
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts) if text_parts else "[No text extracted from PDF]"
    except ImportError:
        pass
    
    try:
        # Fallback ke PyPDF2
        import PyPDF2
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts) if text_parts else "[No text extracted from PDF]"
    except ImportError:
        return "[PDF extraction requires pdfplumber or PyPDF2: pip install pdfplumber PyPDF2]"


def extract_docx_text(file_path: Path) -> str:
    """Extract text dari DOCX menggunakan python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_parts = []
        
        # Extract dari paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract dari tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts) if text_parts else "[No text extracted from DOCX]"
    except ImportError:
        return "[DOCX extraction requires python-docx: pip install python-docx]"


def extract_txt_text(file_path: Path) -> str:
    """Extract text dari file TXT/MD."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Coba dengan encoding lain
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def clean_text(text: str, max_length: int = 50000) -> str:
    """
    Clean dan truncate text untuk storage.
    
    Args:
        text: Raw text
        max_length: Maximum length
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    
    # Truncate jika terlalu panjang
    if len(text) > max_length:
        text = text[:max_length] + f"\n\n[...truncated, total: {len(text)} chars]"
    
    return text.strip()
