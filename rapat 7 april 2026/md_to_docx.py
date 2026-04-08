#!/usr/bin/env python3
"""md_to_docx.py

Skrip konversi Markdown ke DOCX dengan format profesional.

Fitur:
- Membaca file .md dan menghasilkan .docx dengan 100% teks sumber.
- Mengatur ukuran halaman A4, margin normal (2.54 cm), dan font Arial.
- Menjaga struktur heading (H1‑H6) sebagai style heading di DOCX.
- Menambahkan paragraf biasa untuk teks biasa.
- Dapat dipanggil via CLI:
    python md_to_docx.py input.md output.docx
"""

import argparse
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def set_page_layout(doc: Document):
    """Set A4 page size and normal margins (2.54 cm)."""
    sections = doc.sections
    for section in sections:
        # A4 size: 21 cm x 29.7 cm
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        # Normal margins (1 inch ≈ 2.54 cm)
        margin = Cm(2.54)
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def set_default_font(doc: Document, font_name: str = "Arial", font_size_pt: int = 11):
    """Set default font for the whole document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    # For compatibility with older Word versions, set both name and latin
    font.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(font_size_pt)


def add_heading(doc: Document, text: str, level: int):
    """Add heading with appropriate style level (1‑6)."""
    heading = doc.add_heading(text, level=level - 1)  # python-docx uses 0‑4
    # Ensure heading uses Arial as well
    for run in heading.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12 + (6 - level) * 2)  # simple size scaling


def add_paragraph(doc: Document, text: str):
    """Add a normal paragraph preserving line breaks."""
    para = doc.add_paragraph(text)
    para.style = doc.styles["Normal"]
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def parse_markdown(md_text: str, doc: Document):
    """Very simple Markdown parser for headings and paragraphs.
    This keeps the original text 100 %.
    """
    lines = md_text.splitlines()
    buffer = []
    for line in lines:
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            # Flush any buffered paragraph lines first
            if buffer:
                add_paragraph(doc, "\n".join(buffer))
                buffer = []
            hashes, title = heading_match.groups()
            level = len(hashes)
            add_heading(doc, title.strip(), level)
        else:
            buffer.append(line)
    # Add remaining paragraph
    if buffer:
        add_paragraph(doc, "\n".join(buffer))


def main():
    parser = argparse.ArgumentParser(description="Konversi Markdown ke DOCX dengan format profesional.")
    parser.add_argument("input_md", type=Path, help="Path ke file Markdown sumber")
    parser.add_argument("output_docx", type=Path, help="Path output file DOCX")
    args = parser.parse_args()

    if not args.input_md.is_file():
        raise FileNotFoundError(f"File markdown tidak ditemukan: {args.input_md}")

    md_content = args.input_md.read_text(encoding="utf-8")
    doc = Document()
    set_page_layout(doc)
    set_default_font(doc, "Arial", 11)
    parse_markdown(md_content, doc)
    doc.save(args.output_docx)
    print(f"Berhasil mengonversi {args.input_md} → {args.output_docx}")


if __name__ == "__main__":
    main()
