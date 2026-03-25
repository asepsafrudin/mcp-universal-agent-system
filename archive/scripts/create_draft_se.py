#!/usr/bin/env python3
"""
Script untuk membuat draft Surat Edaran Menteri Dalam Negeri
tentang Dukungan Pelaksanaan Sensus Ekonomi 2026
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_draft_se():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Kepala Surat Edaran
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MENTERI DALAM NEGERI\nREPUBLIK INDONESIA")
    run.bold = True
    run.font.size = Pt(12)
    
    # Spasi
    doc.add_paragraph()
    
    # Salinan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SALINAN")
    run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Yth
    p = doc.add_paragraph()
    run = p.add_run("Yth.\n")
    run.font.size = Pt(11)
    run = p.add_run("1. Gubernur selaku Kepala Daerah Provinsi;\n")
    run.font.size = Pt(11)
    run = p.add_run("2. Bupati/Wali Kota selaku Kepala Daerah Kabupaten/Kota;\n")
    run.font.size = Pt(11)
    run = p.add_run("di -\nseluruh Indonesia")
    run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Nomor dan Tentang
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Hide table borders
    for row in table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().clear()
    
    # Row 1 - Nomor
    table.cell(0, 0).text = "Nomor"
    table.cell(0, 1).text = ""
    
    # Row 2 - Tentang
    table.cell(1, 0).text = "Tentang"
    table.cell(1, 1).text = "DUKUNGAN PELAKSANAAN SENSUS EKONOMI 2026"
    
    # Set font size for table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                paragraph.paragraph_format.space_after = Pt(0)
    
    # Make Tentang bold
    for run in table.cell(1, 1).paragraphs[0].runs:
        run.bold = True
    
    # Spasi
    doc.add_paragraph()
    
    # Judul SURAT EDARAN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT EDARAN")
    run.bold = True
    run.font.size = Pt(12)
    
    # Spasi
    doc.add_paragraph()
    
    # Pembukaan
    pembukaan = (
        "Dalam rangka mendukung pelaksanaan Sensus Ekonomi 2026 yang merupakan agenda prioritas "
        "Pemerintah dan berskala nasional, perlu dilakukan koordinasi dan sinergi antara Pemerintah Pusat, "
        "Pemerintah Daerah, dan seluruh pemangku kepentingan terkait untuk memastikan kelancaran, "
        "keberhasilan, dan optimalisasi hasil Sensus Ekonomi 2026."
    )
    p = doc.add_paragraph(pembukaan)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Considerans
    considerans = (
        "Sehubungan dengan hal tersebut, diminta perhatian Gubernur dan Bupati/Wali Kota terhadap "
        "hal-hal sebagai berikut:"
    )
    p = doc.add_paragraph(considerans)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Isi Surat Edaran - 10 poin
    isi_items = [
        "Mendorong dan memastikan partisipasi aktif seluruh unit kerja di lingkungan Pemerintah Daerah dalam mendukung pelaksanaan Sensus Ekonomi 2026 yang akan dilaksanakan pada bulan Mei s.d Juli 2026.",
        
        "Melakukan koordinasi bersama Forum Koordinasi Pimpinan Daerah (Forkopimda), Badan Pusat Statistik (BPS) Provinsi/Kabupaten/Kota, dan seluruh pemangku kepentingan terkait untuk mengidentifikasi potensi hambatan dan menyiapkan langkah-langkah mitigasi dalam pelaksanaan Sensus Ekonomi 2026.",
        
        "Membentuk Tim Koordinasi Sensus Ekonomi 2026 di tingkat Provinsi dan Kabupaten/Kota yang melibatkan unsur terkait, antara lain Bappeda, Dinas Kependudukan dan Catatan Sipil, Dinas Komunikasi dan Informatika, Dinas Penanaman Modal dan Pelayanan Terpadu Satu Pintu, serta instansi terkait lainnya untuk melakukan sinergi, fasilitasi, pengendalian, dan pemantauan pelaksanaan Sensus Ekonomi 2026.",
        
        "Meningkatkan sosialisasi dan edukasi kepada masyarakat dan pelaku usaha tentang pentingnya Sensus Ekonomi 2026 sebagai acuan perencanaan pembangunan nasional dan daerah, serta memastikan partisipasi aktif dalam memberikan data yang akurat dan lengkap.",
        
        "Memastikan ketersediaan data dan informasi pendukung yang diperlukan oleh Badan Pusat Statistik dalam pelaksanaan Sensus Ekonomi 2026, antara lain:",
        
        "Memfasilitasi akses petugas sensus kepada pelaku usaha dan lokasi kegiatan ekonomi di wilayahnya, termasuk koordinasi dengan aparat keamanan setempat untuk memastikan keselamatan dan keamanan petugas sensus selama pelaksanaan kegiatan.",
        
        "Mengoptimalkan pemanfaatan teknologi informasi dan komunikasi dalam pendataan Sensus Ekonomi 2026, antara lain dengan:",
        
        "Menyiapkan langkah-langkah mitigasi terhadap potensi hambatan dalam pelaksanaan Sensus Ekonomi 2026, antara lain:",
        
        "Melakukan monitoring dan evaluasi berkala terhadap pelaksanaan Sensus Ekonomi 2026 di wilayahnya, serta menyampaikan laporan perkembangan secara berjenjang kepada Menteri Dalam Negeri melalui Direktur Jenderal Bina Pembangunan Daerah.",
        
        "Melaporkan hasil pelaksanaan Sensus Ekonomi 2026 dan tindak lanjut dukungan Pemerintah Daerah secara berjenjang kepada Menteri Dalam Negeri melalui Direktur Jenderal Bina Pembangunan Daerah paling lambat 30 (tiga puluh) hari setelah pelaksanaan Sensus Ekonomi 2026 selesai."
    ]
    
    for i, item in enumerate(isi_items, 1):
        if i == 5:  # Poin 5 dengan sub-poin
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            for run in p.runs:
                run.font.size = Pt(11)
            
            # Sub-poin 5
            sub_items_5 = [
                "Data izin usaha dan legalitas pelaku usaha dari Dinas Penanaman Modal dan Pelayanan Terpadu Satu Pintu;",
                "Data pokok kependudukan dari Dinas Kependudukan dan Catatan Sipil;",
                "Data perizinan dan regulasi daerah yang terkait dengan kegiatan ekonomi; dan",
                "Data lainnya yang relevan sesuai dengan kebutuhan Sensus Ekonomi 2026."
            ]
            for sub_item in sub_items_5:
                p = doc.add_paragraph(sub_item)
                p.paragraph_format.left_indent = Inches(0.75)
                for run in p.runs:
                    run.font.size = Pt(11)
        
        elif i == 7:  # Poin 7 dengan sub-poin
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            for run in p.runs:
                run.font.size = Pt(11)
            
            # Sub-poin 7
            sub_items_7 = [
                "Memastikan ketersediaan infrastruktur jaringan internet dan telekomunikasi yang memadai di wilayahnya;",
                "Mendukung penggunaan aplikasi elektronik (CAPI/CAWI) dalam pengumpulan data sensus; dan",
                "Mengoptimalkan layanan call center dan media sosial resmi Pemerintah Daerah untuk menyediakan informasi terkait Sensus Ekonomi 2026."
            ]
            for sub_item in sub_items_7:
                p = doc.add_paragraph(sub_item)
                p.paragraph_format.left_indent = Inches(0.75)
                for run in p.runs:
                    run.font.size = Pt(11)
        
        elif i == 8:  # Poin 8 dengan sub-poin
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            for run in p.runs:
                run.font.size = Pt(11)
            
            # Sub-poin 8
            sub_items_8 = [
                "Mengidentifikasi daerah-daerah yang berpotensi sulit dijangkau dan menyiapkan solusi alternatif;",
                "Menyusun rencana kontingensi untuk mengatasi kondisi darurat atau bencana yang mungkin terjadi selama pelaksanaan sensus; dan",
                "Menyiagakan Tim Reaksi Cepat (TRC) untuk menangani keluhan dan permasalahan yang muncul selama pelaksanaan Sensus Ekonomi 2026."
            ]
            for sub_item in sub_items_8:
                p = doc.add_paragraph(sub_item)
                p.paragraph_format.left_indent = Inches(0.75)
                for run in p.runs:
                    run.font.size = Pt(11)
        
        else:
            p = doc.add_paragraph(f"{i}. {item}")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            for run in p.runs:
                run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Penutup
    penutup = "Demikian untuk menjadi perhatian dan dilaksanakan sebagaimana mestinya."
    p = doc.add_paragraph(penutup)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Spasi untuk tanda tangan
    for _ in range(3):
        doc.add_paragraph()
    
    # Ditetapkan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ditetapkan di Jakarta\n")
    run.font.size = Pt(11)
    run = p.add_run("pada tanggal ...................... 2026")
    run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Jabatan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MENTERI DALAM NEGERI\nREPUBLIK INDONESIA,")
    run.font.size = Pt(11)
    
    # Spasi untuk tanda tangan
    for _ in range(4):
        doc.add_paragraph()
    
    # Nama Menteri
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TITO KARNAVIAN")
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    
    # Page break untuk tembusan
    doc.add_page_break()
    
    # Tembusan
    p = doc.add_paragraph()
    run = p.add_run("Tembusan Yth.:")
    run.font.size = Pt(11)
    
    tembusan_list = [
        "1. Presiden Republik Indonesia;",
        "2. Wakil Presiden Republik Indonesia;",
        "3. Menteri Koordinator Bidang Politik dan Keamanan;",
        "4. Menteri Koordinator Bidang Pembangunan Manusia dan Kebudayaan;",
        "5. Menteri Koordinator Bidang Perekonomian;",
        "6. Menteri Koordinator Bidang Infrastruktur dan Pengembangan Kewilayahan;",
        "7. Menteri Koordinator Bidang Pangan;",
        "8. Menteri Sekretaris Negara;",
        "9. Kepala Badan Pusat Statistik;",
        "10. Menteri Dalam Negeri;",
        "11. Menteri Komunikasi dan Digital;",
        "12. Kepala Kepolisian Negara Republik Indonesia;",
        "13. Kepala Staf Kepresidenan;",
        "14. Sekretaris Kabinet;",
        "15. Ketua Dewan Perwakilan Rakyat Daerah Provinsi seluruh Indonesia; dan",
        "16. Ketua Dewan Perwakilan Rakyat Daerah Kabupaten/Kota seluruh Indonesia."
    ]
    
    for item in tembusan_list:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.size = Pt(10)
    
    # Spasi
    for _ in range(2):
        doc.add_paragraph()
    
    # Salinan sesuai aslinya
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Salinan sesuai dengan aslinya")
    run.font.size = Pt(10)
    
    p = doc.add_paragraph()
    run = p.add_run("Kepala Biro Hukum,")
    run.font.size = Pt(10)
    
    # Spasi
    for _ in range(3):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("NIP.")
    run.font.size = Pt(10)
    
    # Save document
    output_path = "google_drive/SE dukungan Sensus Ekonomi/draft.docx"
    doc.save(output_path)
    print(f"Draft SE berhasil dibuat: {output_path}")

if __name__ == "__main__":
    create_draft_se()
