#!/usr/bin/env python3
"""
Script untuk membuat draft Surat Edaran Menteri Dalam Negeri
tentang Dukungan Pelaksanaan Sensus Ekonomi 2026
Berdasarkan hasil OCR dari file draft SE yang diupload
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_draft_se():
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Lampiran Surat Kepala BPS
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Lampiran Surat Kepala Badan Pusat Statistik\n")
    run.font.size = Pt(10)
    run = p.add_run("Nomor : B-53/01000/SS.190/2026\n")
    run.font.size = Pt(10)
    run = p.add_run("Tanggal : 24 Januari 2026")
    run.font.size = Pt(10)
    
    # Spasi
    doc.add_paragraph()
    
    # Kepala Surat Edaran
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MENTERI DALAM NEGERI\nREPUBLIK INDONESIA")
    run.bold = True
    run.font.size = Pt(12)
    
    # Spasi
    doc.add_paragraph()
    
    # Jakarta dan tanggal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Jakarta, ................... 2026")
    run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Yth
    p = doc.add_paragraph()
    run = p.add_run("1. Sdr/i. Gubernur.\n")
    run.font.size = Pt(11)
    run = p.add_run("2. Sdr/i. Bupati/Walikota.\n")
    run.font.size = Pt(11)
    run = p.add_run("di -\n")
    run.font.size = Pt(11)
    run = p.add_run("Seluruh Indonesia")
    run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Nomor dan Tentang
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Hide table borders
    for row in table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().clear()
    
    # Row 1 - Nomor
    table.cell(0, 0).text = "Nomor"
    table.cell(0, 1).text = ""
    
    # Row 2 - Tentang
    table.cell(1, 0).text = "Tentang"
    table.cell(1, 1).text = "DUKUNGAN PELAKSANAAN SENSUS EKONOMI 2026"
    
    # Set font size for table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                paragraph.paragraph_format.space_after = Pt(0)
    
    # Make Tentang bold
    for run in table.cell(1, 1).paragraphs[0].runs:
        run.bold = True
    
    # Spasi
    doc.add_paragraph()
    
    # Judul SURAT EDARAN
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SURAT EDARAN")
    run.bold = True
    run.font.size = Pt(12)
    
    # Spasi
    doc.add_paragraph()
    
    # Pembukaan
    pembukaan = (
        "Berdasarkan Undang-Undang Nomor 16 Tahun 1997 tentang Statistik dan Peraturan "
        "Pemerintah Nomor 51 Tahun 1999 tentang Penyelenggaraan Statistik, dan Surat Kepala Badan "
        "Pusat Statistik Nomor: B-53/01000/SS.190/2026, tanggal 24 Januari 2026, perihal Permohonan "
        "Dukungan Pelaksanaan Sensus Ekonomi 2026, disampaikan beberapa hal penting sebagai berikut:"
    )
    p = doc.add_paragraph(pembukaan)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Isi Surat Edaran - Poin 1
    p = doc.add_paragraph(
        "1. Pada bulan Mei s.d. Juli 2026, Badan Pusat Statistik (BPS) akan melaksanakan Sensus "
        "Ekonomi 2026 (SE2026) yang merupakan agenda prioritas pemerintah dan berskala "
        "nasional. SE2026 meliputi kegiatan pendaftaran dan pencacahan lengkap seluruh usaha "
        "yang bergerak di semua sektor (kecuali sektor pertanian karena sudah dicakup dalam "
        "Sensus Pertanian 2023)."
    )
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Isi Surat Edaran - Poin 2
    p = doc.add_paragraph(
        "2. Sebagai bentuk peran aktif dalam mendukung kegiatan SE2026, diminta perhatian Gubernur "
        "dan Bupati/Walikota beserta seluruh jajaran di wilayahnya untuk memberi dukungan pada "
        "pelaksanaan SE2026 dalam bentuk berikut ini:"
    )
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Sub-poin a-f
    sub_poins = [
        "membantu mensosialisasikan pelaksanaan SE2026 kepada jajarannya, asosiasi pelaku usaha dan masyarakat di wilayahnya;",
        "mengajak seluruh masyarakat, khususnya para pelaku usaha di wilayahnya untuk berpartisipasi dalam pendataan SE2026 dengan memberikan informasi yang akurat;",
        "memberikan testimoni dukungan dari Gubernur dan Bupati/Walikota pada pelaksanaan kegiatan SE2026 untuk disebarluaskan kepada jajaran pemerintah daerah dan pelaku usaha di wilayah masing-masing. Template narasi testimoni dapat diunduh pada tautan http://s.bps.go.id/NarasiTestimoniSE2026;",
        "membantu menyebarluaskan informasi tentang SE2026 melalui kanal yang dimiliki oleh pemerintah daerah tingkat Provinsi dan Kabupaten/Kota. Termasuk memberikan ruang kepada BPS Provinsi dan Kabupaten/Kota untuk melakukan sosialisasi SE2026 pada kegiatan-kegiatan pemerintah daerah. Materi publisitas SE2026 dapat diakses pada tautan berikut: https://s.bps.go.id/materipublisitasSE2026KLD;",
        "menjalin koordinasi dan konsolidasi terkait pelaksanaan SE2026 dengan jajaran BPS Provinsi dan BPS Kabupaten/Kota di wilayah masing-masing; dan",
        "memfasilitasi keterlibatan Pegawai Pemerintah dengan Perjanjian Kerja (PPPK) di lingkungan Pemerintah Daerah Provinsi dan Kabupaten/Kota sebagai petugas SE2026."
    ]
    
    for i, item in enumerate(sub_poins, ord('a')):
        p = doc.add_paragraph(f"{chr(i)}. {item}")
        p.paragraph_format.left_indent = Inches(0.75)
        for run in p.runs:
            run.font.size = Pt(11)
    
    # Spasi
    doc.add_paragraph()
    
    # Penutup
    penutup = "Demikian untuk menjadi perhatian dan dilaksanakan."
    p = doc.add_paragraph(penutup)
    p.paragraph_format.first_line_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    
    # Spasi untuk tanda tangan
    for _ in range(3):
        doc.add_paragraph()
    
    # Jabatan
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Menteri Dalam Negeri")
    run.font.size = Pt(11)
    
    # Spasi untuk tanda tangan
    for _ in range(3):
        doc.add_paragraph()
    
    # Nama Menteri
    p = doc.add_paragraph()
    run = p.add_run("Prof. H. Muhammad Tito Karnavian, Ph.D")
    run.bold = True
    run.underline = True
    run.font.size = Pt(11)
    
    # Save document
    output_path = "/home/aseps/MCP/google_drive/SE dukungan Sensus Ekonomi/draft.docx"
    doc.save(output_path)
    print(f"✅ Draft SE berhasil dibuat: {output_path}")

if __name__ == "__main__":
    create_draft_se()
